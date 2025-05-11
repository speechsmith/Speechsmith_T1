# import streamlit as st
# from PIL import Image
# from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
# from transformers import VitsTokenizer, VitsModel
# import soundfile as sf
# from openai import OpenAI
# from moviepy.editor import VideoFileClip
# import scipy.io.wavfile as wavfile
# import os
# import io
# import base64
# from gtts import gTTS
# import docx
# import PyPDF2
# import tempfile
# from dotenv import load_dotenv
# import torch
# import numpy as np 
# import json
# import time
# import re
# from deepgram import (
#     DeepgramClient,
#     SpeakOptions,
#     PrerecordedOptions
# )
# from speech_analysis import SpeechAnalyzer, generate_feedback, format_feedback_to_html
# from pydub import AudioSegment
# from groq import Groq
# import subprocess
# import google.generativeai as genai

# load_dotenv()

# # Create necessary directories if they don't exist
# os.makedirs('processed_data', exist_ok=True)
# os.makedirs('processed_data/audio', exist_ok=True)
# os.makedirs('processed_data/text', exist_ok=True)

# openai_api_key = os.getenv("OPENAI_API_KEY")
# hugging_face_token = os.getenv("HUGGINGFACE_TOKEN")
# deepgram_api_key = os.getenv("DEEPGRAM_API_KEY")
# deepgram_client = DeepgramClient(deepgram_api_key)
# gemini_api_key = os.getenv("GEMINI_API_KEY")
# genai.configure(api_key=gemini_api_key)

# def load_services_css():
#     st.markdown("""
#         <style>
#         /* Modern background with dark overlay */
#         .stApp {
#             background: linear-gradient(rgba(0, 0, 0, 0.5), rgba(255, 255, 255, 0.2)),
#                         url("https://github.com/user-attachments/assets/9bc19a87-c89d-405e-94ca-7ad06a920e90") no-repeat center center fixed;
#             background-size: cover;
#         }
        
#         /* Modern gradient text */
#         .gradient-text {
#             background: linear-gradient(135deg, #6C63FF, #FF6B9B);
#             -webkit-background-clip: text;
#             background-clip: text;
#             -webkit-text-fill-color: transparent;
#             font-size: 3.5rem;
#             font-weight: 700;
#             text-align: center;
#             margin-bottom: 2rem;
#             letter-spacing: -0.02em;
#         }
#         .stStatusContainer {
#             font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
#             border-radius: 8px;
#             padding: 1.2rem;
#             margin: 1rem 0;
#             box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
#             transition: all 0.3s ease;
#         }

#         /* Info Message Styling */
#         .stInfo {
#             background: linear-gradient(135deg, #f0f7ff 0%, #e6f3ff 100%);
#             border-left: 4px solid #3b82f6;
#             color: #1e40af;
#         }

#         /* Success Message Styling */
#         .stSuccess {
#             background: linear-gradient(135deg, #f0fdf4 0%, #dcfce7 100%);
#             border-left: 4px solid #22c55e;
#             color: #166534;
#         }

#         .stSuccess::before {
#             content: "✅";
#             margin-right: 0.8rem;
#             font-size: 1.1rem;
#         }

#         /* Error Message Styling */
#         .stError {
#             background: linear-gradient(135deg, #fef2f2 0%, #fee2e2 100%);
#             border-left: 4px solid #ef4444;
#             color: #991b1b;
#             animation: shake 0.5s ease-in-out;
#         }

#         .stError::before {
#             content: "❌";
#             margin-right: 0.8rem;
#             font-size: 1.1rem;
#         }

#         /* Progress Steps Styling */
#         .step-counter {
#             display: inline-block;
#             width: 24px;
#             height: 24px;
#             border-radius: 50%;
#             background-color: #3b82f6;
#             color: white;
#             text-align: center;
#             line-height: 24px;
#             margin-right: 0.8rem;
#             font-size: 0.9rem;
#             font-weight: 500;
#         }

#         /* Animation for Error Messages */
#         @keyframes shake {
#             0%, 100% { transform: translateX(0); }
#             25% { transform: translateX(-4px); }
#             75% { transform: translateX(4px); }
#         }

#         /* Loading Progress Bar */
#         .stProgress {
#             height: 4px;
#             background: linear-gradient(90deg, 
#                 #3b82f6 0%,
#                 #8b5cf6 50%,
#                 #3b82f6 100%
#             );
#             background-size: 200% 100%;
#             animation: progress-animation 2s linear infinite;
#             border-radius: 2px;
#         }

#         @keyframes progress-animation {
#             0% { background-position: 200% 0; }
#             100% { background-position: -200% 0; }
#         }
        
#         /* Modern service cards */
#         .service-card {
#             background: rgba(255, 255, 255, 0.1);
#             backdrop-filter: blur(12px);
#             padding: 2rem;
#             border-radius: 16px;
#             margin-bottom: 1.5rem;
#             height: 220px;
#             display: flex;
#             flex-direction: column;
#             transition: transform 0.3s ease, box-shadow 0.3s ease;
#             border: 1px solid rgba(255, 255, 255, 0.1);
#         }
        
#         .service-card:hover {
#             transform: translateY(-5px);
#             box-shadow: 0 20px 40px rgba(0, 0, 0, 0.2);
#         }
        
#         .service-card h3 {
#             color: #fff;
#             font-size: 1.5rem;
#             margin-bottom: 1rem;
#             font-weight: 600;
#         }
        
#         .service-card p {
#             color: rgba(255, 255, 255, 0.8);
#             font-size: 1.1rem;
#             line-height: 1.6;
#         }
        
#         /* Service icons */
#         .service-icon {
#             color: #6C63FF;
#             font-size: 2.5rem;
#             margin-bottom: 1.5rem;
#             opacity: 0.9;
#         }
                
#         .content-section h3 {
#             font-size: 1.5rem;
#             margin-top: 1.5rem;
#             margin-bottom: 1rem;
#         }

#         /* Paragraph and list styling */
#         .content-section li {
#             color: #334155;
#             font-size: 1.1rem;
#             line-height: 1.75;
#         }

#         /* Card-like section styling */
#         .content-section {
#             background: rgba(255, 255, 255, 0.8);
#             backdrop-filter: blur(12px);
#             border-radius: 16px;
#             padding: 2rem;
#             margin-bottom: 2rem;
#             box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1),
#                         0 2px 4px -1px rgba(0, 0, 0, 0.06);
#         }
                
#         /* Paragraph and list styling */
#         .content-section p, .content-section li {
#             color: #334155;
#             font-size: 1.1rem;
#             line-height: 1.75;
#         }
        
#         .section-subtitle {
#             color: rgba(255, 255, 255, 0.9);
#             font-size: 1.3rem;
#             text-align: center;
#             margin-bottom: 3rem;
#             max-width: 800px;
#             margin-left: auto;
#             margin-right: auto;
#         }
        
#         /* List styling */
#         ul {
#             color: rgba(255, 255, 255, 0.85);
#             margin-left: 1.5rem;
#             margin-bottom: 2rem;
#         }
        
#         li {
#             margin-bottom: 0.75rem;
#             line-height: 1.6;
#         }
        
#         /* Strong text */
#         strong {
#             color: #800080;
#             font-weight: 600;
#         }
        
#         /* Styling for Streamlit error messages */
#         .stError {
#             background: linear-gradient(135deg, #ff6b6b, #ff3b3b);
#             color: #ffffff !important;
#             font-weight: 600;
#             font-size: 1rem;
#             padding: 1rem 1.5rem;
#             border-radius: 8px;
#             box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15);
#             border: 1px solid rgba(255, 0, 0, 0.3);
#             margin-top: 1rem;
#             margin-bottom: 1.5rem;
#             display: flex;
#             align-items: center;
#         }

#         /* Error icon styling for emphasis */
#         .stError::before {
#             content: "⚠️";
#             margin-right: 0.75rem;
#             font-size: 1.25rem;
#         }
        
#         /* Ensure any error message content respects the styling */
#         .stError p {
#             color: #ffffff !important;
#             margin: 0;
#         }
                
#         .generate-speech {
#             display: inline-block;
#             background: rgba(255, 255, 255, 0.8);
#             backdrop-filter: blur(12px);
#             border-radius: 16px;
#             padding: 0.5rem 1rem;
#             margin-bottom: 1rem;
#             color: #334155;
#             box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
#             border: 1px solid rgba(255, 255, 255, 0.2);
#         }
        
#         .stMultiSelect, .stSelectbox, .stTextInput {
#             background: rgba(255, 255, 255, 0.8);
#             backdrop-filter: blur(12px);
#             border-radius: 16px;
#             padding: 1rem;
#             margin-bottom: 1.5rem;
#             color: #334155;
#             box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
#             border: 1px solid rgba(255, 255, 255, 0.2);
#         }
#         .stTextArea {
#             background: rgba(255, 255, 255, 0.8);
#             backdrop-filter: blur(12px);
#             border-radius: 16px;
#             padding: 1rem;
#             margin-bottom: 1.5rem;
#             color: #334155;
#             box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
#             border: 1px solid rgba(255, 255, 255, 0.2);
#         }
        
#         /* File uploader */
#         .stFileUploader {
#             background: rgba(255, 255, 255, 0.7);
#             padding: 2rem;
#             border-radius: 16px;
#             border: 2px dashed rgba(255, 255, 255, 0.5);
#             margin-bottom: 2rem;
#         }
        
#         /* Feedback box */
#         .feedback-box {
#             background: rgba(255, 255, 255, 0.1);
#             backdrop-filter: blur(12px);
#             padding: 2rem;
#             border-radius: 16px;
#             border: 1px solid rgba(255, 255, 255, 0.1);
#             margin-top: 1.5rem;
#             color: white;
#         }
        
#         /* Audio player */
#         .audio-player {
#             width: 100%;
#             margin: 1.5rem 0;
#             background: rgba(255, 255, 255, 0.1);
#             border-radius: 12px;
#             padding: 1rem;
#         }
        
#         /* Divider */
#         hr {
#             border-color: rgba(255, 255, 255, 0.1);
#             margin: 3rem 0;
#         }
        
#         /* Transcription Display Styles */
#         .transcription-container {
#             background: rgba(255, 255, 255, 0.9);
#             padding: 2rem;
#             border-radius: 16px;
#             margin-bottom: 2rem;
#         }
        
#         .transcription-header {
#             margin-bottom: 1rem;
#             padding-bottom: 0.5rem;
#             border-bottom: 2px solid #FF4B4B;
#         }
        
#         .transcription-legend {
#             background: #f8f9fa;
#             padding: 1rem;
#             border-radius: 8px;
#             margin-bottom: 1rem;
#             font-size: 0.9rem;
#         }
        
#         .transcription-text {
#             line-height: 1.8;
#             font-size: 1.1rem;
#         }
        
#         .bold-word {
#             font-weight: bold;
#             color: #FF4B4B;
#         }
        
#         .pause-marker {
#             color: #6C63FF;
#             font-weight: bold;
#         }
        
#         .mispronounced {
#             background-color: rgba(255, 0, 0, 0.1);
#             padding: 0 2px;
#             border-radius: 2px;
#         }
        
#         .play-button {
#             display: inline-block;
#             background: #FF4B4B;
#             color: white;
#             padding: 0.3rem0.8rem;
#             border-radius: 20px;
#             cursor: pointer;
#             margin-right: 0.5rem;
#             font-size: 0.8rem;
#         }
        
#         .play-button:hover {
#             background: #FF3333;
#         }
        
#         .word-container {
#             display: inline-flex;
#             align-items: center;
#             gap: 0.3rem;
#         }
        
#         .play-button {
#             display: inline-flex;
#             align-items: center;
#             justify-content: center;
#             background: #FF4B4B;
#             color: white;
#             width: 24px;
#             height: 24px;
#             border-radius: 50%;
#             border: none;
#             cursor: pointer;
#             padding: 0;
#             font-size: 12px;
#             transition: background-color 0.3s;
#         }
        
#         .play-button:hover {
#             background: #FF3333;
#         }
        
#         .play-button i {
#             margin: 0;
#         }
        
#         /* Add Font Awesome for icons */
#         <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/5.15.4/css/all.min.css">
#         </style>
#     """, unsafe_allow_html=True)

# def extract_text_from_document(uploaded_file):
#     if uploaded_file.name.endswith('.pdf'):
#         pdf_reader = PyPDF2.PdfReader(uploaded_file)
#         text = ""
#         for page in pdf_reader.pages:
#             text += page.extract_text()
#         return text
#     elif uploaded_file.name.endswith('.docx'):
#         doc = docx.Document(uploaded_file)
#         return "\n".join([paragraph.text for paragraph in doc.paragraphs])
#     return None

# def transcribe_audio(audio_file):
#     """Transcribe audio file to text using Gemini"""
#     try:
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
#             temp_audio.write(audio_file.read())
#             temp_audio_path = temp_audio.name
        
#         uploaded = genai.upload_file(temp_audio_path)
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(["Please transcribe the following audio:", uploaded])
#         transcription_text = response.text
#         os.remove(temp_audio_path)
#         return transcription_text
#     except Exception as e:
#         st.error(f"Error in transcription: {str(e)}")
#         return None

# def detect_pauses_with_gemini(audio_file, duration):
#     """Detect pauses in audio using Gemini 1.5 Flash and assess their appropriateness"""
#     try:
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
#             temp_audio.write(audio_file.read())
#             temp_audio_path = temp_audio.name
        
#         uploaded = genai.upload_file(temp_audio_path)
#         model = genai.GenerativeModel('gemini-1.5-flash')
        
#         prompt = """
#         You are an audio analysis expert. Analyze the provided audio file to detect pauses, defined as silent segments lasting at least 0.5 seconds. Provide the following:
#         1. The total number of pauses detected.
#         2. The timestamps (in seconds) where each pause occurs (start time of the silent segment).
        
#         Format the response as:
#         Number of pauses: [count]
#         Pause timestamps: [list of start times in seconds, e.g., 2.5, 7.8, 15.2]
        
#         If no pauses are detected, return:
#         Number of pauses: 0
#         Pause timestamps: []
#         """
        
#         response = model.generate_content([prompt, uploaded])
#         result = response.text.strip()
        
#         os.remove(temp_audio_path)
        
#         pause_count = 0
#         pause_timestamps = []
#         if "Number of pauses:" in result:
#             count_match = re.search(r'Number of pauses: (\d+)', result)
#             if count_match:
#                 pause_count = int(count_match.group(1))
#             timestamps_match = re.search(r'Pause timestamps: \[(.*?)\]', result)
#             if timestamps_match and timestamps_match.group(1):
#                 pause_timestamps = [float(t.strip()) for t in timestamps_match.group(1).split(',') if t.strip()]
        
#         if duration == "Less than 1 minute":
#             if pause_count < 2:
#                 assessment = "Too few, sounding rushed"
#             else:
#                 assessment = "Too many for the given duration"
#         elif duration == "1-3 minutes":
#             if pause_count > 3:
#                 assessment = "Too many for the given duration"
#             else:
#                 assessment = "Too few, sounding rushed"
#         elif duration == "3-5 minutes":
#             if pause_count > 5:
#                 assessment = "Too many for the given duration"
#             else:
#                 assessment = "Too few, sounding rushed"
#         elif duration == "More than 5 minutes":
#             if pause_count > 8:
#                 assessment = "Too many for the given duration"
#             else:
#                 assessment = "Too few, sounding rushed"
#         else:
#             assessment = "Not analyzed"
        
#         return {
#             'count': pause_count,
#             'assessment': assessment,
#             'timestamps': pause_timestamps
#         }
    
#     except Exception as e:
#         st.error(f"Error detecting pauses: {str(e)}")
#         return {
#             'count': 0,
#             'assessment': "Not analyzed",
#             'timestamps': []
#         }


# def calculate_wpm_with_gemini(audio_file, transcription, duration):
#     """Calculate words per minute and assess ideal word count using Gemini 1.5 Flash"""
#     temp_audio_path = None
#     try:
#         # Create temporary file
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
#             temp_audio.write(audio_file.read())
#             temp_audio_path = temp_audio.name
        
#         # Upload to Gemini
#         uploaded = genai.upload_file(temp_audio_path)
#         model = genai.GenerativeModel('gemini-1.5-flash')
        
#         prompt = f"""
#         You are an audio analysis expert. Analyze the provided audio file and its transcription to calculate the speaking rate in words per minute (WPM). Follow these steps:
#         1. Count the total number of words in the transcription.
#         2. Determine the total duration of the audio in seconds.
#         3. Calculate WPM as (total words / duration in minutes), where duration in minutes = duration in seconds / 60.
        
#         Transcription:
#         {transcription}
        
#         Provide the response in the following format:
#         Total words: [count]
#         Audio duration: [seconds] seconds
#         Speaking rate: [wpm] words per minute
#         """
        
#         response = model.generate_content([prompt, uploaded])
#         result = response.text.strip()
        
#         total_words = 0
#         audio_duration = 0.0
#         wpm = 0.0
        
#         # Parse response
#         if "Total words:" in result:
#             words_match = re.search(r'Total words: (\d+)', result)
#             if words_match:
#                 total_words = int(words_match.group(1))
        
#         if "Audio duration:" in result:
#             duration_match = re.search(r'Audio duration: ([\d.]+) seconds', result)
#             if duration_match:
#                 audio_duration = float(duration_match.group(1))
        
#         if "Speaking rate:" in result:
#             wpm_match = re.search(r'Speaking rate: ([\d.]+) words per minute', result)
#             if wpm_match:
#                 wpm = float(wpm_match.group(1))
        
#         # Fallback if parsing fails
#         if total_words == 0 or audio_duration == 0.0 or wpm == 0.0:
#             total_words = len(transcription.split())
#             audio = AudioSegment.from_file(temp_audio_path)
#             audio_duration = len(audio) / 1000.0
#             wpm = (total_words / (audio_duration / 60.0)) if audio_duration > 0 else 0.0
        
#         # Calculate ideal word count and assessment
#         ideal_word_count = ""
#         assessment = ""
#         suggestion = ""
        
#         if duration == "Less than 1 minute":
#             ideal_min, ideal_max = 70, 140
#             if total_words < ideal_min:
#                 assessment = "Too less"
#                 suggestion = f"The current word count is {total_words} words. For a <1 minute speech, ideal count is {ideal_min}-{ideal_max} words - please increase."
#             elif total_words > ideal_max:
#                 assessment = "Too much"
#                 suggestion = f"The current word count is {total_words} words. For a <1 minute speech, ideal count is {ideal_min}-{ideal_max} words - please decrease."
#             else:
#                 assessment = "Appropriate"
#                 suggestion = f"The current word count is {total_words} words, which is appropriate for a <1 minute speech."
        
#         # Add similar logic for other durations (omitted for brevity)
        
#         return {
#             'wpm': round(wpm, 1),
#             'total_words': total_words,
#             'audio_duration': audio_duration,
#             'word_count_assessment': assessment,
#             'word_count_suggestion': suggestion
#         }
    
#     except Exception as e:
#         st.error(f"Error calculating WPM: {str(e)}")
#         try:
#             # Fallback: Calculate WPM manually
#             total_words = len(transcription.split())
#             audio = AudioSegment.from_file(temp_audio_path)
#             audio_duration = len(audio) / 1000.0
#             wpm = (total_words / (audio_duration / 60.0)) if audio_duration > 0 else 0.0
#             return {
#                 'wpm': round(wpm, 1),
#                 'total_words': total_words,
#                 'audio_duration': audio_duration,
#                 'word_count_assessment': 'Not analyzed',
#                 'word_count_suggestion': 'Error in analysis.'
#             }
#         except Exception as inner_e:
#             st.error(f"Fallback WPM calculation failed: {str(inner_e)}")
#             return {
#                 'wpm': 0.0,
#                 'total_words': 0,
#                 'audio_duration': 0.0,
#                 'word_count_assessment': 'Not analyzed',
#                 'word_count_suggestion': 'Error in analysis.'
#             }
    
#     finally:
#         # Clean up temporary file
#         if temp_audio_path and os.path.exists(temp_audio_path):
#             try:
#                 os.remove(temp_audio_path)
#             except Exception as e:
#                 st.warning(f"Failed to delete temporary file {temp_audio_path}: {str(e)}")

# def analyze_speech_with_gemini(transcription, topic=None, duration=None, audio_file=None, gender=None):
#     """Analyze speech using Gemini, with pause detection and WPM calculation"""
#     try:
#         filler_word_list = [
#             "um", "uh", "like", "you know", "so", "actually", "basically", 
#             "seriously", "literally", "well", "okay", "right", "I mean", 
#             "sort of", "stuff", "things", "anyway", "hmm", "ah", 
#             "er", "mm-hmm", "uh-huh", "just", "maybe", "probably", "guess"
#         ]

#         analysis_prompt = f"""You are a professional speech analyst. Your task is to analyze the following speech and provide specific, numerical feedback, including the location of filler words in the input and suggestions for limiting filler words.

# Transcription:
# {transcription}

# Topic: {topic or 'Not specified'}
# Gender: {gender or 'Not specified'}

# IMPORTANT: You MUST provide actual numerical values and specific feedback. Do NOT use placeholders or empty values.

# Format your analysis as follows:

# Pronunciation
# Overall speech accuracy: [percentage]%
# Feedback: [detailed feedback on pronunciation clarity and issues]

# Mood
# Primary emotion: [emotion, e.g., Confidence]
# Formality level: [e.g., Professional]
# Audience suitability: [e.g., Highly suitable for academic audience]
# Assessment: [evaluation of emotional tone and suitability]
# Reasons:
# - [Reason 1, e.g., Consistent use of professional terminology]
# - [Reason 2, e.g., Measured pace shows confidence]
# - [Reason 3, e.g., Clear emphasis on key points]

# Speaking Style
# Filler Words Analysis:
# - [filler word, e.g., "um"]: [number] occurrences
# - [filler word, e.g., "like"]: [number] occurrences
# Total filler words: [total count]
# Filler word percentage: [percentage]%
# Location:
#   [e.g., Beginning: "Hi, um, my name is Anushka Tandan"]
#   [e.g., After "Tandan": "um and I would like"]
# Assessment: [evaluation of filler word usage]
# Suggestions for Limiting Filler Words:
# 1. [Suggestion 1, e.g., Pause Intentionally: Instead of saying "um," simply pause.]
# 2. [Suggestion 2, e.g., Plan Your Opening]
# 3. [Suggestion 3, e.g., Record Yourself and Listen Back]
# 4. [Suggestion 4, e.g., Practice Concise Speaking]

# Pitch
# Pitch Analysis:
# Pitch variation: [e.g., Good variation, Excessive variation, or Too monotonous]
# Consistency: [e.g., Very consistent with appropriate emphasis]
# Average pitch: [Qualitative assessment, e.g., Too high for the given gender, Too low for the given gender, Appropriate for the given gender]

# Example:
# Pronunciation
# Overall speech accuracy: 85%
# Feedback: Clear enunciation with good articulation. Some difficulty with complex words like "statistics" and "phenomenon".

# Mood
# Primary emotion: Confidence
# Formality level: Professional
# Audience suitability: Highly suitable for academic audience
# Assessment: The emotional tone effectively conveys authority and expertise
# Reasons:
# - Consistent use of professional terminology
# - Measured pace shows confidence
# - Clear emphasis on key points

# Speaking Style
# Filler Words Analysis:
# - "um": 3 occurrences
# - "like": 2 occurrences
# Total filler words: 2
# Filler word percentage: 2.1%
# Location:
#   Beginning: "Hi, um, my name is Anushka Tandan"
#   After "Tandan": "um and I would like"
# Assessment: Moderate use of filler words disrupts flow
# Suggestions for Limiting Filler Words:
# 1. Pause Intentionally: Instead of saying "um," simply pause.
# 2. Plan Your Opening: A well-rehearsed opening reduces filler words.
# 3. Record Yourself: Listen to identify filler word patterns.
# 4. Practice Concise Speaking: Be direct in communication.

# Pitch
# Pitch Analysis:
# Pitch variation: Good variation
# Consistency: Very consistent with appropriate emphasis
# Average pitch: Appropriate for the given gender

# By implementing these strategies, you can significantly reduce your use of filler words and communicate with more clarity and confidence. Do not copy the examples - provide your own analysis based on the actual speech content.
# """

#         model = genai.GenerativeModel('gemini-1.5-flash')
#         analysis_response = model.generate_content(analysis_prompt)
#         analysis_text = analysis_response.text
        
#         accuracy_match = re.search(r'Overall speech accuracy: (\d+)%', analysis_text)
#         accuracy = int(accuracy_match.group(1)) if accuracy_match else 0
        
#         # Update pitch parsing to handle qualitative average pitch
#         pitch_variation_match = re.search(r'Pitch variation: (.+?)\n', analysis_text)
#         pitch_variation = pitch_variation_match.group(1).strip() if pitch_variation_match else 'Not analyzed'
        
#         consistency_match = re.search(r'Consistency: (.+?)\n', analysis_text)
#         consistency = consistency_match.group(1).strip() if consistency_match else 'Not analyzed'
        
#         average_pitch_match = re.search(r'Average pitch: (.+?)(?:\n|$)', analysis_text)
#         average_pitch = average_pitch_match.group(1).strip() if average_pitch_match else 'Not analyzed'

#         filler_words = {word: 0 for word in filler_word_list}
#         transcription_lower = transcription.lower()
#         for word in filler_word_list:
#             pattern = r'\b' + re.escape(word) + r'\b'
#             count = len(re.findall(pattern, transcription_lower))
#             filler_words[word] = count
        
#         filler_words = {k: v for k, v in filler_words.items() if v > 0}
#         total_fillers = sum(filler_words.values())
#         total_words = len(transcription.split())
#         filler_percentage = (total_fillers / total_words * 100) if total_words > 0 else 0
        
#         filler_locations = []
#         filler_suggestions = []
#         filler_assessment = ''
#         if 'Filler Words Analysis:' in analysis_text:
#             filler_section = analysis_text.split('Filler Words Analysis:')[1].split('Pitch')[0]
#             location_section = filler_section.split('Location:')[1].split('Assessment:')[0].strip() if 'Location:' in filler_section else ''
#             for line in location_section.split('\n'):
#                 line = line.strip()
#                 if line:
#                     filler_locations.append(line)
            
#             filler_assessment = filler_section.split('Assessment:')[1].split('Suggestions')[0].strip() if 'Assessment:' in filler_section else ''
            
#             suggestions_section = filler_section.split('Suggestions for Limiting Filler Words:')[1].strip() if 'Suggestions for Limiting Filler Words:' in filler_section else ''
#             for line in suggestions_section.split('\n'):
#                 line = line.strip()
#                 if line and not line.startswith('By implementing these strategies'):
#                     filler_suggestions.append(line)
        
#         mood_section = analysis_text.split('Mood')[1].split('Speaking Style')[0] if 'Mood' in analysis_text and 'Speaking Style' in analysis_text else ''
#         primary_emotion = mood_section.split('Primary emotion:')[1].split('\n')[0].strip() if 'Primary emotion:' in mood_section else ''
#         formality = mood_section.split('Formality level:')[1].split('\n')[0].strip() if 'Formality level:' in mood_section else ''
#         audience_suitability = mood_section.split('Audience suitability:')[1].split('\n')[0].strip() if 'Audience suitability:' in mood_section else ''
#         mood_assessment = mood_section.split('Assessment:')[1].split('Reasons:')[0].strip() if 'Assessment:' in mood_section and 'Reasons:' in mood_section else ''
        
#         reasons = []
#         if 'Reasons:' in mood_section:
#             reasons_section = mood_section.split('Reasons:')[1].strip()
#             for line in reasons_section.split('\n'):
#                 line = line.strip()
#                 if line.startswith('-') and line.strip('-').strip():
#                     reasons.append(line.strip('-').strip())
        
#         wpm_data = {'wpm': 0.0, 'total_words': 0, 'audio_duration': 0.0, 'word_count_assessment': 'Not analyzed', 'word_count_suggestion': 'Not analyzed'}
#         if audio_file and duration:
#             audio_file.seek(0)
#             wpm_data = calculate_wpm_with_gemini(audio_file, transcription, duration)
        
#         pause_data = {'count': 0, 'assessment': 'Not analyzed', 'timestamps': []}
#         if audio_file and duration:
#             audio_file.seek(0)
#             pause_data = detect_pauses_with_gemini(audio_file, duration)
        
#         analysis_dict = {
#             'pronunciation': {
#                 'accuracy': accuracy,
#                 'feedback': analysis_text.split('Feedback:')[1].split('\n\n')[0].strip() if 'Feedback:' in analysis_text else '',
#                 'difficult_words': {}
#             },
#             'pitch': {
#                 'variation': pitch_variation,
#                 'consistency': consistency,
#                 'average': average_pitch  # Now a qualitative string
#             },
#             'speech_rate': {
#                 'wpm': wpm_data['wpm'],
#                 'total_words': wpm_data['total_words'],
#                 'word_count_assessment': wpm_data['word_count_assessment'],
#                 'word_count_suggestion': wpm_data['word_count_suggestion'],
#                 'pauses': {
#                     'count': pause_data['count'],
#                     'assessment': pause_data['assessment']
#                 },
#                 'filler_words': filler_words,
#                 'filler_word_percentage': filler_percentage,
#                 'total_filler_words': total_fillers,
#                 'filler_locations': filler_locations,
#                 'filler_assessment': filler_assessment,
#                 'filler_suggestions': filler_suggestions
#             },
#             'mood': {
#                 'primary_emotion': primary_emotion,
#                 'formality': formality,
#                 'audience_suitability': audience_suitability,
#                 'mood_suitability_assessment': {
#                     'assessment': mood_assessment,
#                     'reasons': reasons
#                 }
#             }
#         }
        
#         analysis_dict['raw_analysis'] = analysis_text
#         return analysis_dict
            
#     except Exception as e:
#         st.error(f"Error in speech analysis: {str(e)}")
#         return None

# def process_with_gemini(transcription, purpose, audience, duration, tone, additional_requirements, topic, speech_analysis):
#     """Generate detailed feedback using Gemini"""
#     try:
#         if not transcription or not transcription.strip():
#             return "No transcription provided", "", "Please provide a valid transcription."

#         transcription = transcription.strip()
#         mispronounced_words = []
#         if speech_analysis and isinstance(speech_analysis, dict):
#             pronunciation = speech_analysis.get('pronunciation', {})
#             if isinstance(pronunciation, dict):
#                 mispronounced_words = pronunciation.get('difficult_words', [])
        
#         system_prompt = """You are a professional speech coach and analyst. 
#         Your task is to provide detailed feedback on the speech delivery and content.
#         Focus on making the feedback comprehensive and actionable.
        
#         Format your feedback in the following sections:
#         1. Content Analysis
#         2. Delivery Analysis
#         3. Specific Recommendations
#         4. Overall Assessment
#         """
        
#         user_prompt = f"""
#         Please provide detailed feedback on this speech:

#         Original Text: {transcription}
#         Topic: {topic or 'Not specified'}
#         Purpose: {purpose or 'Not specified'}
#         Target Audience: {audience or 'Not specified'}
#         Duration: {duration or 'Not specified'}
#         Tone: {tone or 'Not specified'}
#         Additional Requirements: {additional_requirements or 'None'}
#         Mispronounced Words: {mispronounced_words}

#         Provide comprehensive feedback covering:
#         1. Content Analysis:
#            - Message clarity
#            - Structure and organization
#            - Relevance to topic and audience
#            - Use of supporting evidence/examples
        
#         2. Delivery Analysis:
#            - Speaking pace and rhythm
#            - Voice modulation and emphasis
#            - Pronunciation and articulation
#            - Use of pauses and transitions
        
#         3. Specific Recommendations:
#            - Areas for improvement
#            - Specific techniques to try
#            - Content adjustments
#            - Delivery enhancements
        
#         4. Overall Assessment:
#            - Strengths
#            - Areas for growth
#            - Next steps for improvement
#         """
        
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content([system_prompt, user_prompt])
#         feedback = response.text
#         return transcription, transcription, feedback

#     except Exception as e:
#         return transcription, transcription, f"Error in generating feedback: {str(e)}"


# def identify_mispronounced_words(audio_file, transcription):
#     """Identify mispronounced words from an audio file using Gemini and return feedback."""
#     import re
#     import streamlit as st
#     import tempfile
#     import os
#     from google.generativeai import GenerativeModel, upload_file
    
#     try:
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
#             temp_audio.write(audio_file.read())
#             temp_audio_path = temp_audio.name
        
#         uploaded = upload_file(temp_audio_path)
#         model = GenerativeModel('gemini-1.5-flash')
        
#         prompt = """
#         You are a pronunciation expert analyzing a speech audio file and its transcription. Your task is to:
#         1. Evaluate the pronunciation accuracy of the speaker based on standard American English.
#         2. Identify all mispronounced words, providing details for each, but only if the errors are clear and significant.

#         Transcription:
#         {transcription}

#         For each mispronounced word, provide:
#         - Word: The mispronounced word.
#         - Correct Pronunciation: A clear, syllable-by-syllable breakdown (e.g., YOU-ni-VER-si-tee).
#         - Phonetic Transcription: The IPA transcription (e.g., /ˌjuːnɪˈvɜːrsəti/).
#         - Explanation: Why the word was mispronounced and how to correct it.

#         Format your response as follows:
#         Pronunciation Feedback: [Overall assessment of pronunciation accuracy]
#         Mispronounced Words:
#         - Word: [word]
#           Correct Pronunciation: [syllable breakdown]
#           Phonetic Transcription: [IPA]
#           Explanation: [details]
#         [If no mispronounced words are detected, state: "No mispronounced words detected."]
#         Additional Notes: [Any additional observations or limitations, e.g., audio quality, need for visual confirmation, or acceptable regional variations]

#         Guidelines:
#         - Only identify words as mispronounced if there are clear and significant errors that impact clarity (e.g., incorrect stress, slurred syllables, wrong vowel sounds, or substitutions that change the word's meaning).
#         - Do not flag minor variations, such as slight accent differences, acceptable regional pronunciations, fast speech rate, or subtle intonation differences, unless they significantly affect intelligibility.
#         - If the pronunciation is clear and accurate with no significant errors, return "No mispronounced words detected" without listing any words. This is the expected outcome for accurate speech.
#         - Use standard American English pronunciation as the reference unless otherwise specified.
#         - Consider potential limitations, such as audio quality, background noise, or speech rate, and note these in the Additional Notes section. Do not flag issues caused by these limitations as mispronunciations unless they cause significant errors.
#         - Avoid overanalyzing or assuming errors without strong evidence from the audio or transcription.
#         """
        
#         prompt = prompt.format(transcription=transcription)
#         response = model.generate_content([prompt, uploaded])
#         result = response.text.strip()
        
#         os.remove(temp_audio_path)
        
#         # Debug: Log transcription and raw response
#         st.write("Transcription:", transcription)
#         st.write("Raw Gemini Pronunciation Response:", result)
        
#         pronunciation_feedback = ""
#         mispronounced_words = []
        
#         # Extract feedback
#         feedback_match = re.search(r'Pronunciation Feedback: (.+?)(?:\nMispronounced Words:|$)', result, re.DOTALL)
#         if feedback_match:
#             pronunciation_feedback = feedback_match.group(1).strip()
        
#         # Check for "No mispronounced words detected"
#         if "No mispronounced words detected" in result:
#             mispronounced_words = []
#         else:
#             # Extract mispronounced words section
#             mispronounced_section = re.search(r'Mispronounced Words:\n(.+?)(?:\nAdditional Notes:|\n\n|$)', result, re.DOTALL)
#             if mispronounced_section:
#                 section_text = mispronounced_section.group(1).strip()
                
#                 # Try multi-line format first
#                 word_blocks = re.split(r'\n- Word: ', section_text)[1:]  # Skip first empty split
#                 for block in word_blocks:
#                     word_match = re.match(
#                         r'(.+?)\n\s{2}Correct Pronunciation: (.+?)\n\s{2}Phonetic Transcription: (.+?)\n\s{2}Explanation: (.+?)(?=\n- Word:|\nAdditional Notes:|\n\n|$)',
#                         block.strip(), re.DOTALL
#                     )
#                     if word_match:
#                         word, correct_pron, phonetic, explanation = word_match.groups()
#                         mispronounced_words.append({
#                             'word': word.strip(),
#                             'correct_pronunciation': correct_pron.strip(),
#                             'phonetic_transcription': phonetic.strip(),
#                             'explanation': explanation.strip()
#                         })
                
#                 # If no matches, try single-line format
#                 if not mispronounced_words:
#                     word_matches = re.findall(
#                         r'- Word: (.+?) Correct Pronunciation: (.+?) Phonetic Transcription: (.+?) Explanation: (.+?)(?=\s*- Word:|$|\n)',
#                         section_text, re.DOTALL
#                     )
#                     for word, correct_pron, phonetic, explanation in word_matches:
#                         mispronounced_words.append({
#                             'word': word.strip(),
#                             'correct_pronunciation': correct_pron.strip(),
#                             'phonetic_transcription': phonetic.strip(),
#                             'explanation': explanation.strip()
#                         })
        
#         # Debug: Log parsed mispronounced words
#         st.write("Parsed Mispronounced Words:", mispronounced_words)
        
#         # Extract additional notes
#         notes_match = re.search(r'Additional Notes: (.+?)(?:\n\n|$)', result, re.DOTALL)
#         notes = notes_match.group(1).strip() if notes_match else ""
        
#         return {
#             'mispronounced_words': mispronounced_words,
#             'pronunciation_feedback': pronunciation_feedback,
#             'notes': notes
#         }
    
#     except Exception as e:
#         st.error(f"Error identifying mispronounced words: {str(e)}")
#         return {
#             'mispronounced_words': [],
#             'pronunciation_feedback': 'Analysis failed due to an error.',
#             'notes': ''
#         }
  
# def process_with_gpt(openai_api_key, transcription, purpose, audience, duration, tone, additional_requirements, topic, speech_analysis):
#     """Process transcription and generate refined speech using GPT, synchronizing bold markers."""
#     try:
#         if not transcription or not transcription.strip():
#             return "No transcription provided", "", "Please provide a valid transcription."

#         transcription = transcription.strip()
#         mispronounced_words = []
#         if speech_analysis and isinstance(speech_analysis, dict):
#             pronunciation = speech_analysis.get('pronunciation', {})
#             if isinstance(pronunciation, dict):
#                 mispronounced_words = pronunciation.get('difficult_words', [])

#         system_prompt = """You are a professional speech writer and coach. 
#         Your task is to refine the given speech and provide detailed feedback on its delivery.
#         Focus on making the speech more engaging and appropriate for the specified audience and purpose.
        
#         For the transcription and refined speech:
#         1. Add **bold** around words that should be emphasized
#         2. Add | (vertical bar) where the speaker should pause
#         3. Mark mispronounced words with <mispronounced> tags
        
#         Format your response exactly as follows:
#         ORIGINAL:
#         [transcription with markers]
        
#         REFINED:
#         [refined speech with markers]
        
#         FEEDBACK:
#         [detailed feedback]
#         """

#         user_prompt = f"""
#         Please analyze and refine this speech:

#         Original Text: {transcription}
#         Topic: {topic or 'Not specified'}
#         Purpose: {purpose or 'Not specified'}
#         Target Audience: {audience or 'Not specified'}
#         Duration: {duration or 'Not specified'}
#         Tone: {tone or 'Not specified'}
#         Mispronounced Words: {mispronounced_words}

#         Provide:
#         1. The original transcription with emphasis markers and pause indicators
#         2. A refined version of the speech with the same markers
#         3. Detailed feedback on delivery and content
#         """

#         groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
#         completion = groq_client.chat.completions.create(
#             messages=[
#                 {"role": "system", "content": system_prompt},
#                 {"role": "user", "content": user_prompt}
#             ],
#             model="llama-3.3-70b-versatile",
#             temperature=0.1,
#             max_tokens=1000,
#         )

#         full_response = completion.choices[0].message.content
#         parts = full_response.split('\n\n')
#         original = ""
#         refined = ""
#         feedback = ""

#         for part in parts:
#             if part.startswith('ORIGINAL:'):
#                 original = part.replace('ORIGINAL:', '').strip()
#             elif part.startswith('REFINED:'):
#                 refined = part.replace('REFINED:', '').strip()
#             elif part.startswith('FEEDBACK:'):
#                 feedback = part.replace('FEEDBACK:', '').strip()

#         if not original or not refined:
#             original = transcription
#             words = original.split()
#             for i, word in enumerate(words):
#                 if len(word) > 4 and word[0].isupper():
#                     words[i] = f"**{word}**"
#             original = ' '.join(words)
#             return original, transcription, "Could not process the response properly. Using original transcription with basic emphasis."

#         # Extract bolded words from refined speech
#         bolded_words = []
#         bold_pattern = r'\*\*(.*?)\*\*'
#         for match in re.finditer(bold_pattern, refined):
#             bolded_word = match.group(1).strip()
#             bolded_words.append(bolded_word.lower())

#         # Apply bold markers to original transcription
#         def apply_bold_to_word(word):
#             if word.lower() in bolded_words:
#                 return f"**{word}**"
#             return word

#         # Split original transcription, preserving pauses and spaces
#         parts = re.split(r'(\s+|\|)', original)
#         modified_parts = []
#         for part in parts:
#             if part.strip() and part not in ['|', ' ']:
#                 modified_parts.append(apply_bold_to_word(part))
#             else:
#                 modified_parts.append(part)
        
#         modified_original = ''.join(modified_parts)

#         return modified_original, refined, feedback

#     except Exception as e:
#         return transcription, transcription, f"Error in processing: {str(e)}"

# def generate_audio_from_text(client, text, filename="temp_output.wav"):
#     try:
#         if not text or not text.strip():
#             st.error("No text provided for audio generation")
#             return None

#         speak_options = {'text': text.strip()}
#         options = SpeakOptions(model="aura-asteria-en", encoding="linear16", container="wav")
#         response = client.speak.v("1").save(filename, speak_options, options)
        
#         with open(filename, 'rb') as audio_file:
#             audio_bytes = audio_file.read()
#             audio_base64 = base64.b64encode(audio_bytes).decode('utf-8')
        
#         if os.path.exists(filename):
#             os.remove(filename)
            
#         return audio_base64
#     except Exception as e:
#         st.error(f"Error generating audio: {e}")
#         return None

# def save_processed_data(session_id, data_type, content):
#     """Save processed data to local storage"""
#     timestamp = time.strftime("%Y%m%d-%H%M%S")
#     if data_type == 'text':
#         filepath = f'processed_data/text/{session_id}_{timestamp}.txt'
#         with open(filepath, 'w') as f:
#             f.write(content)
#     elif data_type == 'audio':
#         filepath = f'processed_data/audio/{session_id}_{timestamp}.wav'
#         with open(filepath, 'wb') as f:
#             f.write(content)
#     return filepath

# def load_processed_data(filepath):
#     """Load processed data from local storage"""
#     if not os.path.exists(filepath):
#         return None
    
#     if filepath.endswith('.txt'):
#         with open(filepath, 'r') as f:
#             return f.read()
#     elif filepath.endswith('.wav'):
#         with open(filepath, 'rb') as f:
#             return f.read()
#     return None

# def convert_mp3_to_wav(mp3_file):
#     """Convert MP3 file to WAV format"""
#     try:
#         audio = AudioSegment.from_mp3(mp3_file)
#         wav_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
#         audio.export(wav_file.name, format="wav")
#         return wav_file.name
#     except Exception as e:
#         st.error(f"Error converting MP3 to WAV: {str(e)}")
#         return None

# def convert_mp4_to_wav(mp4_file):
#     """Convert MP4 file to WAV format"""
#     try:
#         if not mp4_file:
#             st.error("No file provided for conversion")
#             return None

#         with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp_mp4:
#             tmp_mp4.write(mp4_file.getvalue())
#             mp4_path = tmp_mp4.name

#         wav_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
#         wav_path = wav_file.name
#         audio = AudioSegment.from_file(mp4_path, format="mp4")
#         audio = audio.set_frame_rate(44100).set_channels(1)
#         audio.export(wav_path, format="wav")
        
#         os.unlink(mp4_path)
#         return wav_path
            
#     except Exception as e:
#         if 'mp4_path' in locals() and os.path.exists(mp4_path):
#             os.unlink(mp4_path)
#         if 'wav_path' in locals() and os.path.exists(wav_path):
#             os.unlink(wav_path)
#         st.error(f"Error converting MP4 to WAV: {str(e)}")
#         return None

# def generate_word_pronunciation(word):
#     """Generate audio for a single word using gTTS"""
#     try:
#         # tts = gTTS(text=word, lang='en')
#         tts = gTTS(text = word, lang='en', tld='co.in')
#         audio_io = io.BytesIO()
#         tts.write_to_fp(audio_io)
#         audio_io.seek(0)
#         return audio_io
#     except Exception as e:
#         st.error(f"Error generating pronunciation for {word}: {str(e)}")
#         return None

# def format_transcription_text(text, mispronounced_words=None):
#     """Format transcription text with emphasis markers and mispronounced words"""
#     if mispronounced_words is None:
#         mispronounced_words = []
    
#     # Convert bold markers to HTML
#     pattern = r'\*\*(.*?)\*\*'
#     text = re.sub(pattern, r'<span class="bold-word">\1</span>', text)
#     # Convert pause markers to HTML
#     text = text.replace('|', '<span class="pause-marker">|</span>')
    
#     # Split text, preserving HTML tags and spaces
#     parts = re.split(r'(\s+|<[^>]*>)', text)
#     formatted_parts = []
    
#     for part in parts:
#         if part.strip() and not part.startswith('<'):
#             word = part.strip()
#             # Remove punctuation for mispronounced word comparison
#             clean_word = re.sub(r'[^\w\s]', '', word).lower()
#             if any(clean_word == mw.lower() for mw in mispronounced_words):
#                 formatted_parts.append(f'<span class="mispronounced">{word}</span>')
#             else:
#                 formatted_parts.append(word)
#         else:
#             formatted_parts.append(part)
    
#     return ''.join(formatted_parts)

# def format_transcription_with_emphasis(transcription, mispronounced_words=None):
#     """Format transcription with emphasis markers and mispronounced words"""
#     if mispronounced_words is None:
#         mispronounced_words = []
    
#     pattern = r'\*\*(.*?)\*\*'
#     transcription = re.sub(pattern, r'<span class="bold-word">\1</span>', transcription)
    
#     parts = re.split(r'(\s+|<[^>]*>)', transcription)
#     formatted_parts = []
    
#     for part in parts:
#         if part.strip() and not part.startswith('<'):
#             word = part.strip()
#             word_lower = word.lower()
#             if word_lower in mispronounced_words:
#                 formatted_parts.append(f'<span class="mispronounced">{word}</span>')
#             else:
#                 formatted_parts.append(word)
#         else:
#             formatted_parts.append(part)
    
#     formatted_text = ''.join(formatted_parts)
    
#     legend = """
#     <div class="transcription-legend">
#         <strong>Legend:</strong><br>
#         <span class="bold-word">Bold words</span> - Words to emphasize<br>
#         <span class="pause-marker">|</span> - Pause in speech<br>
#         <span class="mispronounced">Highlighted words</span> - Words that need pronunciation improvement
#     </div>
#     """
    
#     return f"""
#     <div class="transcription-container">
#         {legend}
#         <div class="transcription-text">
#             {formatted_text}
#         </div>
#     </div>
  
#     """


# def format_detailed_feedback(results, pronunciation_data=None):
#     """Format detailed feedback sections with dynamic data, including pronunciation feedback."""
#     import streamlit as st
    
#     if isinstance(results, str):
#         return f"""
#         <div class="content-section">
#             <h2>Detailed Feedback</h2>
#             <pre>{results}</pre>
#         </div>
#         """
    
#     pronunciation = results.get('pronunciation', {})
#     pitch = results.get('pitch', {})
#     speech_rate = results.get('speech_rate', {})
#     mood = results.get('mood', {})
    
#     # Ensure pronunciation_data is a dictionary
#     if not isinstance(pronunciation_data, dict):
#         pronunciation_data = {
#             'mispronounced_words': [],
#             'pronunciation_feedback': 'Not analyzed (invalid pronunciation data)',
#             'notes': ''
#         }
#         st.error("Invalid pronunciation data format. Expected a dictionary.")
    
#     # Debug: Log pronunciation_data
#     st.write("Pronunciation Data for Rendering:", pronunciation_data)
    
#     pronunciation_section = """
#     <h3>Pronunciation</h3>
#     <ul>
#         <li><strong>Overall speech accuracy:</strong>{accuracy}%</li>
#         <li><strong>Pronunciation Feedback:</strong>{feedback}</li>
#     """.format(
#         accuracy=pronunciation.get('accuracy', 0),
#         feedback=pronunciation_data.get('pronunciation_feedback', 'Not analyzed')
#     )
    
#     mispronounced_words = pronunciation_data.get('mispronounced_words', [])
#     if mispronounced_words:
#         pronunciation_section += "<li><strong>Mispronounced Words:</strong></li><ul>"
#         for item in mispronounced_words:
#             pronunciation_section += f"""
#             <li>
#                 Word: {item['word']}<br>
#                 Correct Pronunciation: {item['correct_pronunciation']}<br>
#                 Phonetic Transcription: {item['phonetic_transcription']}<br>
#                 Explanation: {item['explanation']}
#             </li>
#             """
#         pronunciation_section += "</ul>"
#     else:
#         pronunciation_section += "<li><strong>Mispronounced Words:</strong>No mispronounced words detected.</li>"
    
#     notes = pronunciation_data.get('notes', '')
#     if notes:
#         pronunciation_section += f"<li><strong>Additional Notes:</strong>{notes}</li>"
    
#     pronunciation_section += "</ul>"
    
#     mood_section = """
#     <h3>Mood</h3>
#     <ul>
#         <li><strong>Primary emotion:</strong>{primary_emotion}</li>
#         <li><strong>Formality level:</strong>{formality}</li>
#         <li><strong>Audience suitability:</strong>{audience}</li>
#         <li><strong>Assessment</strong>{assessment}</li>
#     """.format(
#         primary_emotion=mood.get('primary_emotion', 'Not analyzed'),
#         formality=mood.get('formality', 'Not analyzed'),
#         audience=mood.get('audience_suitability', 'Not analyzed'),
#         assessment=mood.get('mood_suitability_assessment', {}).get('assessment', 'Not analyzed')
#     )
    
#     reasons = mood.get('mood_suitability_assessment', {}).get('reasons', [])
#     if reasons:
#         mood_section += "<li><strong>Reasons:</strong></li><ul>"
#         for reason in reasons:
#             mood_section += f"<li>{reason}</li>"
#         mood_section += "</ul>"
#     else:
#         mood_section += "<li><strong>Reasons:</strong> No specific reasons provided.</li>"
    
#     mood_section += "</ul>"
    
#     speaking_style_section = """
#     <h3>Speaking Style</h3>
#     <ul>
#         <li><strong>Speaking rate:</strong> {rate} words per minute</li>
#         <li><strong>Ideal word count:</strong>{word_count_assessment}({word_count_suggestion})</li>
#         <li><strong>Pauses:</strong> {pause_assessment}</li>
#     """.format(
#         rate=speech_rate.get('wpm', 0),
#         word_count_assessment=speech_rate.get('word_count_assessment', 'Not analyzed'),
#         word_count_suggestion=speech_rate.get('word_count_suggestion', 'Not analyzed'),
#         pause_assessment=speech_rate.get('pauses', {}).get('assessment', 'Not analyzed')
#     )
    
#     if speech_rate.get('filler_words'):
#         speaking_style_section += "<li><strong>Filler words analysis:</strong></li><ul>"
#         for word, count in speech_rate['filler_words'].items():
#             speaking_style_section += f"<li>'{word}': {count} occurrences</li>"
#         speaking_style_section += "</ul>"
#     speaking_style_section += f"<li><strong>Total filler words:</strong>{speech_rate.get('total_filler_words', 0)}</li>"
    
#     if speech_rate.get('filler_locations'):
#         speaking_style_section += "<li><strong>Location:</strong></li><ul>"
#         for location in speech_rate['filler_locations']:
#             speaking_style_section += f"<li>{location}</li>"
#         speaking_style_section += "</ul>"
    
#     if speech_rate.get('filler_assessment'):
#         speaking_style_section += f"<li><strong>Assessment:</strong>{speech_rate['filler_assessment']}</li>"
    
#     if speech_rate.get('filler_suggestions'):
#         speaking_style_section += "<li><strong>Suggestions for Limiting Filler Words:</strong></li><ul>"
#         for suggestion in speech_rate['filler_suggestions']:
#             speaking_style_section += f"<li>{suggestion}</li>"
#         speaking_style_section += "</ul>"
#         speaking_style_section += "<li>By implementing these strategies, you can significantly reduce your use of filler words and communicate with more clarity and confidence.</li>"
    
#     speaking_style_section += "</ul>"
    
#     pitch_section = """
#     <h3>Pitch</h3>
#     <ul>
#         <li><strong>Pitch Analysis:</strong></li>
#         <ul>
#             <li>Detected Gender: {detected_gender}</li>
#             <li>Pitch variation: {variation}</li>
#             <li>Consistency: {consistency}</li>
#             <li>Average pitch: {average}</li>
#         </ul>
#     </ul>
#     """.format(
#         detected_gender=pitch.get('detected_gender', 'Not analyzed'),
#         variation=pitch.get('variation', 'Not analyzed'),
#         consistency=pitch.get('consistency', 'Not analyzed'),
#         average=pitch.get('average', 'Not analyzed')
#     )
    
#     # Debug: Log generated HTML
#     html_content = f"""
#     <div class="content-section">
#         <h2>Detailed Feedback</h2>
#         {pronunciation_section}
#         {mood_section}
#         {speaking_style_section}
#         {pitch_section}
#     </div>
#     """
#     st.write("Generated HTML Content:", html_content)
    
#     return html_content

# def analyze_audio_gender_and_pitch(audio_file):
#     """Analyze the gender and pitch of an audio file using Gemini."""
#     try:
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
#             temp_audio.write(audio_file.read())
#             temp_audio_path = temp_audio.name
        
#         uploaded = genai.upload_file(temp_audio_path)
#         model = genai.GenerativeModel('gemini-1.5-flash')
        
#         prompt = """
#         You are an audio analysis expert specializing in voice characteristics. Your task is to analyze the provided audio file to:
#         1. Detect the speaker's gender (male, female, or indeterminate if unclear).
#         2. Evaluate the pitch characteristics, including:
#            - Whether the average pitch is too high, too low, or appropriate for the detected gender.
#            - Whether the pitch variation is excessive, too monotonous, or appropriate.

#         Provide the response in the following format:
#         Detected Gender: [male, female, or indeterminate]
#         Average Pitch: [Too high for the given gender, Too low for the given gender, Appropriate for the given gender, or Cannot be definitively determined without clear gender detection]
#         Pitch Variation: [Excessive, Too monotonous, or Appropriate]

#         Guidelines:
#         - For gender detection, use vocal characteristics such as pitch, timbre, and resonance. Typical ranges are:
#           - Male: 85-180 Hz (lower pitch, deeper resonance).
#           - Female: 165-255 Hz (higher pitch, lighter resonance).
#           - Indeterminate: If characteristics overlap significantly or are unclear.
#         - For average pitch:
#           - Compare the speaker's average pitch to the typical range for the detected gender.
#           - Male: Classify as "Too high" if >190 Hz, "Too low" if <80 Hz, "Appropriate" if 80-190 Hz.
#           - Female: Classify as "Too high" if >260 Hz, "Too low" if <150 Hz, "Appropriate" if 150-260 Hz.
#           - Indeterminate: Use a neutral range (120-220 Hz) for assessment.
#         - For pitch variation:
#           - "Excessive" if pitch fluctuates >80 Hz across the audio, indicating overly dramatic shifts.
#           - "Too monotonous" if pitch variation is <15 Hz, indicating a flat delivery.
#           - "Appropriate" if variation is 15-80 Hz, supporting natural, engaging speech.
#         - If gender is indeterminate, provide a fallback pitch assessment based on the neutral range (120-220 Hz).
#         - Be precise and avoid defaulting to "Appropriate" unless the pitch is clearly within the specified range for the gender.
#         - If the audio is too short (<3 seconds) or unclear, note limitations in the analysis.
#         """
        
#         response = model.generate_content([prompt, uploaded])
#         result = response.text.strip()
        
#         # Debug: Log raw response
#         # st.write("Raw Gemini Response:", result)
        
#         os.remove(temp_audio_path)
        
#         detected_gender = "indeterminate"
#         average_pitch = "Cannot be definitively determined without clear gender detection"
#         pitch_variation = "Not analyzed"
        
#         gender_match = re.search(r'Detected Gender: (male|female|indeterminate)', result)
#         if gender_match:
#             detected_gender = gender_match.group(1)
        
#         pitch_match = re.search(r'Average Pitch: (.+?)(?:\n|$)', result)
#         if pitch_match:
#             average_pitch = pitch_match.group(1).strip()
        
#         variation_match = re.search(r'Pitch Variation: (.+?)(?:\n|$)', result)
#         if variation_match:
#             pitch_variation = variation_match.group(1).strip()
        
#         return {
#             'detected_gender': detected_gender,
#             'average_pitch': average_pitch,
#             'pitch_variation': pitch_variation
#         }
    
#     except Exception as e:
#         st.error(f"Error analyzing audio gender and pitch: {str(e)}")
#         return {
#             'detected_gender': 'indeterminate',
#             'average_pitch': 'Cannot be definitively determined due to analysis error',
#             'pitch_variation': 'Not analyzed'
#         }



# def services():
#     analyzer = SpeechAnalyzer()
    
#     if 'usage_count' not in st.session_state:
#         st.session_state.usage_count = 0
    
#     if st.session_state.usage_count >= 5 and not st.session_state.get('is_authenticated', False):
#         st.error("""
#             You have reached the free trial limit of 5 uses.
#             Please fill out the Contact Us form to request access credentials.
#             Our team will review your request and provide you with login details.
#             Thank you for trying out SpeechSmith!
#         """)
#         return

#     load_services_css()
    
#     if 'session_id' not in st.session_state:
#         st.session_state.session_id = f"session_{time.strftime('%Y%m%d%H%M%S')}"
    
#     if 'text_filepath' not in st.session_state:
#         st.session_state.text_filepath = None
#     if 'audio_filepath' not in st.session_state:
#         st.session_state.audio_filepath = None
    
#     st.markdown('<h1 class="gradient-text">Our Services</h1>', unsafe_allow_html=True)
#     st.markdown('<p class="section-subtitle">At speechsmith, we offer a seamless and effective way to refine your speech delivery, ensuring it meets your specific goals and resonates with your audience.</p>', unsafe_allow_html=True)
    
#     if not st.session_state.get('is_authenticated', False):
#         remaining_uses = 5 - st.session_state.usage_count
#         st.info(f"You have {remaining_uses} free trial uses remaining. Please contact us for full access.")
    
#     col1, col2 = st.columns(2)
    
#     with col1:
#         st.markdown("""
#             <div class="service-card">
#                 <i class="fas fa-microphone service-icon"></i>
#                 <h3>AUDIO UPLOAD</h3>
#                 <p>Easily upload your speech recordings to our platform for comprehensive analysis.</p>
#             </div>
            
#             <div class="service-card">
#                 <i class="fas fa-tasks service-icon"></i>
#                 <h3>CUSTOMIZED SPEECH GOALS</h3>
#                 <p>Tailor your speech refinement by selecting your target audience and the intent of your speech.</p>
#             </div>
#         """, unsafe_allow_html=True)
    
#     with col2:
#         st.markdown("""
#             <div class="service-card">
#                 <i class="fas fa-comments service-icon"></i>
#                 <h3>DETAILED FEEDBACK</h3>
#                 <p>Receive in-depth feedback on your speech, including insights on tone, pacing, and clarity.</p>
#             </div>
            
#             <div class="service-card">
#                 <i class="fas fa-edit service-icon"></i>
#                 <h3>REFORMED SPEECH</h3>
#                 <p>Get a refined version of your speech that aligns perfectly with your chosen audience and intent.</p>
#             </div>
#         """, unsafe_allow_html=True)
    
#     st.markdown("---")
#     st.markdown('<h2 class="gradient-text">Upload Your Speech</h2>', unsafe_allow_html=True)
    
#     uploaded_file = st.file_uploader("Upload your audio file", type=['mp3', 'wav', 'mp4'])
#     uploaded_document = st.file_uploader("Upload the script as a pdf/word document (optional, but suggested)", type=['pdf', 'docx'])
    
#     if uploaded_file:
#         st.subheader("Original Speech")
#         st.audio(uploaded_file, format='audio/wav')
#         st.session_state.original_audio = uploaded_file

#     topic_of_speech = ""
    
#     if uploaded_file is not None or uploaded_document is not None:
#         topic_of_speech = st.text_input("Enter the topic of the speech")
#         gender = st.selectbox("Select Gender", ['male', 'female', 'Other', 'Prefer not to say'])
#         purpose = st.selectbox("What is the purpose of your speech?", ["Inform", "Persuade/Inspire", "Entertain", "Other"])
#         if purpose == "Other":
#             purpose = st.text_input("Please specify the purpose")

#         audience = st.selectbox("Who is your target audience?", ["Classmates/Colleagues", "Teachers/Professors", "General public", "Other"])
#         if audience == "Other":
#             audience = st.text_input("Please specify the audience")

#         duration = st.selectbox("How long is your speech intended to be?", ["Less than 1 minute", "1-3 minutes", "3-5 minutes", "More than 5 minutes"])
#         tone = st.selectbox("What tone do you wish to adopt?", ["Formal", "Informal", "Humorous", "Other"])
#         if tone == "Other":
#             tone = st.text_input("Please specify the tone")
        
#         additional_requirements = st.text_area("Any additional requirements or preferences? (Optional)", height=100)
        
#         if st.button("Process Speech"):
#             if not purpose or not audience or not duration or not tone:
#                 st.error("Please fill out all required fields before processing.")
#                 st.stop()
            
#             if not st.session_state.get('is_authenticated', False):
#                 st.session_state.usage_count += 1
            
#             status_container = st.empty()
            
#             try:
#                 status_container.markdown("""
#                     <div class="stStatusContainer stInfo">
#                         <span class="step-counter">1</span>
#                         Transcribing your speech using Gemini...
#                     </div>
#                 """, unsafe_allow_html=True)
                
#                 if uploaded_file:
#                     transcription = transcribe_audio(uploaded_file)
#                     if not transcription or not transcription.strip():
#                         st.error("No speech detected in the audio file")
#                         st.stop()
                    
#                     st.session_state.transcription = transcription
                    
#                     status_container.markdown("""
#                         <div class="stStatusContainer stInfo">
#                             <span class="step-counter">2</span>
#                             Identifying mispronounced words...
#                         </div>
#                     """, unsafe_allow_html=True)
                    
#                     uploaded_file.seek(0)
#                     pronunciation_data = identify_mispronounced_words(uploaded_file, transcription)
                    
#                     status_container.markdown("""
#                         <div class="stStatusContainer stInfo">
#                             <span class="step-counter">3</span>
#                             Analyzing gender and pitch...
#                         </div>
#                     """, unsafe_allow_html=True)
                    
#                     uploaded_file.seek(0)
#                     pitch_analysis = analyze_audio_gender_and_pitch(uploaded_file)
                    
#                     try:
#                         st.info("Generating AI version of your speech...")
#                         tts = gTTS(text=transcription, lang='en')
#                         ai_audio_io = io.BytesIO()
#                         tts.write_to_fp(ai_audio_io)
#                         ai_audio_io.seek(0)
#                         ai_audio_bytes = ai_audio_io.getvalue()
                        
#                         st.session_state.ai_audio_io = ai_audio_io
#                         st.session_state.ai_audio_bytes = ai_audio_bytes
#                         st.success("AI version generated successfully!")
#                     except Exception as e:
#                         st.error(f"Error generating AI version: {str(e)}")
                    
#                     status_container.markdown("""
#                         <div class="stStatusContainer stInfo">
#                             <span class="step-counter">4</span>
#                             Analyzing speech metrics...
#                         </div>
#                     """, unsafe_allow_html=True)
                    
#                     results = analyze_speech_with_gemini(transcription, topic_of_speech, duration, uploaded_file, gender=pitch_analysis['detected_gender'])
#                     if not results:
#                         st.error("Failed to analyze speech")
#                         st.stop()
                    
#                     # Override pitch analysis with audio-based results
#                     results['pitch'] = {
#                         'variation': pitch_analysis['pitch_variation'],
#                         'consistency': results['pitch'].get('consistency', 'Not analyzed'),
#                         'average': pitch_analysis['average_pitch'],
#                         'detected_gender': pitch_analysis['detected_gender']
#                     }
                    
#                     # Override pronunciation feedback for consistency
#                     results['pronunciation']['feedback'] = pronunciation_data['pronunciation_feedback']
                    
#                 else:
#                     transcription = extract_text_from_document(uploaded_document)
#                     if not transcription or not transcription.strip():
#                         st.error("No text found in the document")
#                         st.stop()
#                     results = analyze_speech_with_gemini(transcription, topic_of_speech, duration, gender=gender)
#                     pronunciation_data = {
#                         'mispronounced_words': [],
#                         'pronunciation_feedback': 'Not analyzed (audio required)',
#                         'notes': ''
#                     }
#                     # For document-only input, use user-selected gender and default pitch analysis
#                     results['pitch'] = {
#                         'variation': 'Not analyzed (audio required)',
#                         'consistency': 'Not analyzed (audio required)',
#                         'average': 'Appropriate for the given gender (cannot be definitively determined without audio)',
#                         'detected_gender': gender
#                     }
                
#                 status_container.markdown("""
#                     <div class="stStatusContainer stInfo">
#                         <span class="step-counter">5</span>
#                         Generating refined speech and detailed feedback...
#                     </div>
#                 """, unsafe_allow_html=True)
                
#                 original, refined, _ = process_with_gpt(
#                     openai_api_key, transcription, purpose, audience, duration, tone, additional_requirements, topic_of_speech, results
#                 )
                
#                 _, _, feedback = process_with_gemini(
#                     transcription, purpose, audience, duration, tone, additional_requirements, topic_of_speech, results
#                 )

#                 status_container.markdown("""
#                     <div class="stStatusContainer stInfo">
#                         <span class="step-counter">6</span>
#                         Generating and saving refined speech...
#                     </div>
#                 """, unsafe_allow_html=True)
                
#                 text_filepath = save_processed_data(st.session_state.session_id, 'text', refined)
#                 st.session_state.text_filepath = text_filepath

#                 status_container.markdown("""
#                     <div class="stStatusContainer stInfo">
#                         <span class="step-counter">7</span>
#                         Generating audio from refined speech...
#                     </div>
#                 """, unsafe_allow_html=True)
                
#                 cleaned_speech = re.sub(r'\*\*|\*_pause_\*|_|#|<[^>]*>', '', refined)
#                 cleaned_speech = cleaned_speech.strip()
                
#                 if not cleaned_speech:
#                     st.error("No text available for audio generation")
#                     st.stop()

#                 tts = gTTS(text=cleaned_speech)
#                 audio_io = io.BytesIO()
#                 tts.write_to_fp(audio_io)
#                 audio_io.seek(0)
#                 audio_base64 = base64.b64encode(audio_io.read()).decode()

#                 if audio_base64:
#                     audio_bytes = base64.b64decode(audio_base64)
#                     audio_filepath = save_processed_data(st.session_state.session_id, 'audio', audio_bytes)
#                     st.session_state.audio_filepath = audio_filepath

#                 status_container.markdown("""
#                     <div class="stStatusContainer stSuccess">
#                         <span class="step-counter">✓</span>
#                         Processing complete! Showing results...
#                     </div>
#                 """, unsafe_allow_html=True)

#                 # Store results in session state for persistence
#                 st.session_state.original = original
#                 st.session_state.refined = refined
#                 st.session_state.results = results
#                 st.session_state.pronunciation_data = pronunciation_data

#                 # Render the results after processing
#                 content = f"""
#                 <div class="content-section">
#                     <h3>Original Transcription</h3>
#                     <div class="transcription-container">
#                         <div class="transcription-legend">
#                             <strong>Legend:</strong><br>
#                             <span class="bold-word">Bold words</span> - Words to emphasize<br>
#                             <span class="pause-marker">|</span> - Pause in speech<br>
#                             <span class="mispronounced">Highlighted words</span> - Words that need pronunciation improvement
#                         </div>
#                         <div class="transcription-text">
#                             {format_transcription_text(original, [item['word'] for item in pronunciation_data['mispronounced_words']])}
#                         </div>
#                     </div>
                    
#                     <h3>Refined Speech</h3>
#                     {format_transcription_with_emphasis(refined, [item['word'] for item in pronunciation_data['mispronounced_words']])}
#                 </div>
#                 {format_detailed_feedback(results, pronunciation_data)}
#                 """
#                 st.html(content)
                
#                 st.markdown("""
#                     <style>
#                         .mispronounced {
#                             background-color: rgba(255, 0, 0, 0.1);
#                             padding: 0 2px;
#                             border-radius: 2px;
#                             color: #ff0000;
#                             font-weight: bold;
#                         }
#                         .transcription-text {
#                             white-space: pre-wrap;
#                             word-wrap: break-word;
#                             font-size: 16px;
#                             line-height: 1.6;
#                             padding: 15px;
#                             background-color: #f8f9fa;
#                             border-radius: 5px;
#                         }
#                     </style>
#                 """, unsafe_allow_html=True)
                
#                 if audio_base64:
#                     content = """
#                         <div class="generate-speech">
#                             <h3>Generated Speech Audio</h3>
#                         </div>
#                         """
#                     st.html(content)
#                     st.audio(io.BytesIO(audio_bytes), format="audio/wav")
                
#                     if st.session_state.get('ai_audio_bytes'):
#                         st.subheader("AI Version of Original Speech")
#                         try:
#                             ai_audio_io = io.BytesIO(st.session_state.ai_audio_bytes)
#                             ai_audio_io.seek(0)
#                             st.audio(ai_audio_io, format='audio/mp3')
                            
#                             download_container = st.container()
                            
#                             with download_container:
#                                 st.markdown(
#                                     f'<a href="data:audio/mp3;base64,{base64.b64encode(st.session_state.ai_audio_bytes).decode()}" download="ai_version.mp3" style="text-decoration: none;">'
#                                     '<button style="background-color: #4CAF50; color: white; padding: 10px 20px; border: none; border-radius: 4px; cursor: pointer; width: 100%;">'
#                                     'Download AI Version'
#                                     '</button>'
#                                     '</a>',
#                                     unsafe_allow_html=True
#                                 )
                            
#                             with download_container:
#                                 st.markdown(
#                                     f'<a href="data:audio/wav;base64,{audio_base64}" download="refined_speech.wav" style="text-decoration: none;">'
#                                     '<button style="background-color: #4CAF50; color: white; padding: 10px 20px; border: none; border-radius: 4px; cursor: pointer; width: 100%;">'
#                                     'Download Refined Speech Audio'
#                                     '</button>'
#                                     '</a>',
#                                     unsafe_allow_html=True
#                                 )
#                         except Exception as e:
#                             st.error(f"Error displaying AI audio: {str(e)}")
#                             st.write("Debug info - Audio data length:", len(st.session_state.ai_audio_bytes) if st.session_state.get('ai_audio_bytes') else "No audio data")
            
#             except Exception as e:
#                 if not st.session_state.get('is_authenticated', False):
#                     st.session_state.usage_count -= 1
#                 status_container.error(f"An error occurred: {str(e)}")
#                 st.stop()
            
#     if st.session_state.get('results'):
#         st.subheader("Analysis Results")
        
#         if uploaded_file:
#             st.subheader("Original Speech")
#             st.audio(uploaded_file, format='audio/wav')
            
#             if st.session_state.get('ai_audio_bytes'):
#                 st.subheader("AI Version of Original Speech")
#                 try:
#                     ai_audio_io = io.BytesIO(st.session_state.ai_audio_bytes)
#                     ai_audio_io.seek(0)
#                     st.audio(ai_audio_io, format='audio/mp3')
#                     st.download_button(
#                         label="Download AI Version",
#                         data=st.session_state.ai_audio_bytes,
#                         file_name="ai_version.mp3",
#                         mime="audio/mp3"
#                     )
#                 except Exception as e:
#                     st.error(f"Error displaying AI audio: {str(e)}")

# def main():
#     st.set_page_config(
#         page_title="SpeechSmith Services",
#         page_icon="🎤",
#         layout="wide"
#     )
#     services()

# if __name__ == "__main__":
#     main()




import streamlit as st
from PIL import Image
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
from transformers import VitsTokenizer, VitsModel
import soundfile as sf
from openai import OpenAI
from moviepy.editor import VideoFileClip
from google.generativeai import GenerativeModel, upload_file
import scipy.io.wavfile as wavfile
import os
import io
import base64
from tempfile import NamedTemporaryFile
from gtts import gTTS
import docx
import PyPDF2
import tempfile
from dotenv import load_dotenv
import torch
import numpy as np 
import json
import time
import re
from deepgram import (
    DeepgramClient,
    SpeakOptions,
    PrerecordedOptions
)
from speech_analysis import SpeechAnalyzer, generate_feedback, format_feedback_to_html
from pydub import AudioSegment
from groq import Groq
import subprocess
import google.generativeai as genai

load_dotenv()

# Create necessary directories if they don't exist
os.makedirs('processed_data', exist_ok=True)
os.makedirs('processed_data/audio', exist_ok=True)
os.makedirs('processed_data/text', exist_ok=True)

openai_api_key = os.getenv("OPENAI_API_KEY")
hugging_face_token = os.getenv("HUGGINGFACE_TOKEN")
deepgram_api_key = os.getenv("DEEPGRAM_API_KEY")
deepgram_client = DeepgramClient(deepgram_api_key)
gemini_api_key = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=gemini_api_key)

def load_services_css():
    st.markdown("""
        <style>
        /* Modern background with dark overlay */
        .stApp {
            background: linear-gradient(rgba(0, 0, 0, 0.5), rgba(255, 255, 255, 0.2)),
                        url("https://github.com/user-attachments/assets/9bc19a87-c89d-405e-94ca-7ad06a920e90") no-repeat center center fixed;
            background-size: cover;
        }
        
        /* Modern gradient text */
        .gradient-text {
            background: linear-gradient(135deg, #6C63FF, #FF6B9B);
            -webkit-background-clip: text;
            background-clip: text;
            -webkit-text-fill-color: transparent;
            font-size: 3.5rem;
            font-weight: 700;
            text-align: center;
            margin-bottom: 2rem;
            letter-spacing: -0.02em;
        }
        .stStatusContainer {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            border-radius: 8px;
            padding: 1.2rem;
            margin: 1rem 0;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            transition: all 0.3s ease;
        }

        /* Info Message Styling */
        .stInfo {
            background: linear-gradient(135deg, #f0f7ff 0%, #e6f3ff 100%);
            border-left: 4px solid #3b82f6;
            color: #1e40af;
        }

        /* Success Message Styling */
        .stSuccess {
            background: linear-gradient(135deg, #f0fdf4 0%, #dcfce7 100%);
            border-left: 4px solid #22c55e;
            color: #166534;
        }

        .stSuccess::before {
            content: "✅";
            margin-right: 0.8rem;
            font-size: 1.1rem;
        }

        /* Error Message Styling */
        .stError {
            background: linear-gradient(135deg, #fef2f2 0%, #fee2e2 100%);
            border-left: 4px solid #ef4444;
            color: #991b1b;
            animation: shake 0.5s ease-in-out;
        }

        .stError::before {
            content: "❌";
            margin-right: 0.8rem;
            font-size: 1.1rem;
        }

        /* Progress Steps Styling */
        .step-counter {
            display: inline-block;
            width: 24px;
            height: 24px;
            border-radius: 50%;
            background-color: #3b82f6;
            color: white;
            text-align: center;
            line-height: 24px;
            margin-right: 0.8rem;
            font-size: 0.9rem;
            font-weight: 500;
        }

        /* Animation for Error Messages */
        @keyframes shake {
            0%, 100% { transform: translateX(0); }
            25% { transform: translateX(-4px); }
            75% { transform: translateX(4px); }
        }

        /* Loading Progress Bar */
        .stProgress {
            height: 4px;
            background: linear-gradient(90deg, 
                #3b82f6 0%,
                #8b5cf6 50%,
                #3b82f6 100%
            );
            background-size: 200% 100%;
            animation: progress-animation 2s linear infinite;
            border-radius: 2px;
        }

        @keyframes progress-animation {
            0% { background-position: 200% 0; }
            100% { background-position: -200% 0; }
        }
        
        /* Modern service cards */
        .service-card {
            background: rgba(255, 255, 255, 0.1);
            backdrop-filter: blur(12px);
            padding: 2rem;
            border-radius: 16px;
            margin-bottom: 1.5rem;
            height: 220px;
            display: flex;
            flex-direction: column;
            transition: transform 0.3s ease, box-shadow 0.3s ease;
            border: 1px solid rgba(255, 255, 255, 0.1);
        }
        
        .service-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 20px 40px rgba(0, 0, 0, 0.2);
        }
        
        .service-card h3 {
            color: #fff;
            font-size: 1.5rem;
            margin-bottom: 1rem;
            font-weight: 600;
        }
        
        .service-card p {
            color: rgba(255, 255, 255, 0.8);
            font-size: 1.1rem;
            line-height: 1.6;
        }
        
        /* Service icons */
        .service-icon {
            color: #6C63FF;
            font-size: 2.5rem;
            margin-bottom: 1.5rem;
            opacity: 0.9;
        }
                
        .content-section h3 {
            font-size: 1.5rem;
            margin-top: 1.5rem;
            margin-bottom: 1rem;
        }

        /* Paragraph and list styling */
        .content-section li {
            color: #334155;
            font-size: 1.1rem;
            line-height: 1.75;
        }

        /* Card-like section styling */
        .content-section {
            background: rgba(255, 255, 255, 0.8);
            backdrop-filter: blur(12px);
            border-radius: 16px;
            padding: 2rem;
            margin-bottom: 2rem;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1),
                        0 2px 4px -1px rgba(0, 0, 0, 0.06);
        }
                
        /* Paragraph and list styling */
        .content-section p, .content-section li {
            color: #334155;
            font-size: 1.1rem;
            line-height: 1.75;
        }
        
        .section-subtitle {
            color: rgba(255, 255, 255, 0.9);
            font-size: 1.3rem;
            text-align: center;
            margin-bottom: 3rem;
            max-width: 800px;
            margin-left: auto;
            margin-right: auto;
        }
        
        /* List styling */
        ul {
            color: rgba(255, 255, 255, 0.85);
            margin-left: 1.5rem;
            margin-bottom: 2rem;
        }
        
        li {
            margin-bottom: 0.75rem;
            line-height: 1.6;
        }
        
        /* Strong text */
        strong {
            color: #800080;
            font-weight: 600;
        }
        
        /* Styling for Streamlit error messages */
        .stError {
            background: linear-gradient(135deg, #ff6b6b, #ff3b3b);
            color: #ffffff !important;
            font-weight: 600;
            font-size: 1rem;
            padding: 1rem 1.5rem;
            border-radius: 8px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15);
            border: 1px solid rgba(255, 0, 0, 0.3);
            margin-top: 1rem;
            margin-bottom: 1.5rem;
            display: flex;
            align-items: center;
        }

        /* Error icon styling for emphasis */
        .stError::before {
            content: "⚠️";
            margin-right: 0.75rem;
            font-size: 1.25rem;
        }
        
        /* Ensure any error message content respects the styling */
        .stError p {
            color: #ffffff !important;
            margin: 0;
        }
                
        .generate-speech {
            display: inline-block;
            background: rgba(255, 255, 255, 0.8);
            backdrop-filter: blur(12px);
            border-radius: 16px;
            padding: 0.5rem 1rem;
            margin-bottom: 1rem;
            color: #334155;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
            border: 1px solid rgba(255, 255, 255, 0.2);
        }
        
        .stMultiSelect, .stSelectbox, .stTextInput {
            background: rgba(255, 255, 255, 0.8);
            backdrop-filter: blur(12px);
            border-radius: 16px;
            padding: 1rem;
            margin-bottom: 1.5rem;
            color: #334155;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
            border: 1px solid rgba(255, 255, 255, 0.2);
        }
        .stTextArea {
            background: rgba(255, 255, 255, 0.8);
            backdrop-filter: blur(12px);
            border-radius: 16px;
            padding: 1rem;
            margin-bottom: 1.5rem;
            color: #334155;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
            border: 1px solid rgba(255, 255, 255, 0.2);
        }
        
        /* File uploader */
        .stFileUploader {
            background: rgba(255, 255, 255, 0.7);
            padding: 2rem;
            border-radius: 16px;
            border: 2px dashed rgba(255, 255, 255, 0.5);
            margin-bottom: 2rem;
        }
        
        /* Feedback box */
        .feedback-box {
            background: rgba(255, 255, 255, 0.1);
            backdrop-filter: blur(12px);
            padding: 2rem;
            border-radius: 16px;
            border: 1px solid rgba(255, 255, 255, 0.1);
            margin-top: 1.5rem;
            color: white;
        }
        
        /* Audio player */
        .audio-player {
            width: 100%;
            margin: 1.5rem 0;
            background: rgba(255, 255, 255, 0.1);
            border-radius: 12px;
            padding: 1rem;
        }
        
        /* Divider */
        hr {
            border-color: rgba(255, 255, 255, 0.1);
            margin: 3rem 0;
        }
        
        /* Transcription Display Styles */
        .transcription-container {
            background: rgba(255, 255, 255, 0.9);
            padding: 2rem;
            border-radius: 16px;
            margin-bottom: 2rem;
        }
        
        .transcription-header {
            margin-bottom: 1rem;
            padding-bottom: 0.5rem;
            border-bottom: 2px solid #FF4B4B;
        }
        
        .transcription-legend {
            background: #f8f9fa;
            padding: 1rem;
            border-radius: 8px;
            margin-bottom: 1rem;
            font-size: 0.9rem;
        }
        
        .transcription-text {
            line-height: 1.8;
            font-size: 1.1rem;
        }
        
        .bold-word {
            font-weight: bold;
            color: #FF4B4B;
        }
        
        .pause-marker {
            color: #6C63FF;
            font-weight: bold;
        }
        
        .mispronounced {
            background-color: rgba(255, 0, 0, 0.1);
            padding: 0 2px;
            border-radius: 2px;
        }
        
        .play-button {
            display: inline-block;
            background: #FF4B4B;
            color: white;
            padding: 0.3rem 0.8rem;
            border-radius: 20px;
            cursor: pointer;
            margin-right: 0.5rem;
            font-size: 0.8rem;
        }
        
        .play-button:hover {
            background: #FF3333;
        }
        
        .word-container {
            display: inline-flex;
            align-items: center;
            gap: 0.3rem;
        }
        
        .play-button {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            background: #FF4B4B;
            color: white;
            width: 24px;
            height: 24px;
            border-radius: 50%;
            border: none;
            cursor: pointer;
            padding: 0;
            font-size: 12px;
            transition: background-color 0.3s;
        }
        
        .play-button:hover {
            background: #FF3333;
        }
        
        .play-button i {
            margin: 0;
        }
        
        /* Add Font Awesome for icons */
        <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/5.15.4/css/all.min.css">
        </style>
    """, unsafe_allow_html=True)


def extract_text_from_document(uploaded_file):
    if uploaded_file.name.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif uploaded_file.name.endswith('.docx'):
        doc = docx.Document(uploaded_file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return None

def transcribe_audio(audio_file):
    """Transcribe audio file to text using Gemini"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
            temp_audio.write(audio_file.read())
            temp_audio_path = temp_audio.name
        
        uploaded = genai.upload_file(temp_audio_path)
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(["Please transcribe the following audio:", uploaded])
        transcription_text = response.text
        os.remove(temp_audio_path)
        return transcription_text
    except Exception as e:
        st.error(f"Error in transcription: {str(e)}")
        return None

def detect_pauses_with_gemini(audio_file, duration):
    """Detect pauses in audio using Gemini 1.5 Flash and assess their appropriateness"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
            temp_audio.write(audio_file.read())
            temp_audio_path = temp_audio.name
        
        uploaded = genai.upload_file(temp_audio_path)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = """
        You are an audio analysis expert. Analyze the provided audio file to detect pauses, defined as silent segments lasting at least 0.5 seconds. Provide the following:
        1. The total number of pauses detected.
        2. The timestamps (in seconds) where each pause occurs (start time of the silent segment).
        
        Format the response as:
        Number of pauses: [count]
        Pause timestamps: [list of start times in seconds, e.g., 2.5, 7.8, 15.2]
        
        If no pauses are detected, return:
        Number of pauses: 0
        Pause timestamps: []
        """
        
        response = model.generate_content([prompt, uploaded])
        result = response.text.strip()
        
        os.remove(temp_audio_path)
        
        pause_count = 0
        pause_timestamps = []
        if "Number of pauses:" in result:
            count_match = re.search(r'Number of pauses: (\d+)', result)
            if count_match:
                pause_count = int(count_match.group(1))
            timestamps_match = re.search(r'Pause timestamps: \[(.*?)\]', result)
            if timestamps_match and timestamps_match.group(1):
                pause_timestamps = [float(t.strip()) for t in timestamps_match.group(1).split(',') if t.strip()]
        
        if duration == "Less than 1 minute":
            if pause_count < 2:
                assessment = "Too few, sounding rushed"
            else:
                assessment = "Too many for the given duration"
        elif duration == "1-3 minutes":
            if pause_count > 3:
                assessment = "Too many for the given duration"
            else:
                assessment = "Too few, sounding rushed"
        elif duration == "3-5 minutes":
            if pause_count > 5:
                assessment = "Too many for the given duration"
            else:
                assessment = "Too few, sounding rushed"
        elif duration == "More than 5 minutes":
            if pause_count > 8:
                assessment = "Too many for the given duration"
            else:
                assessment = "Too few, sounding rushed"
        else:
            assessment = "Not analyzed"
        
        return {
            'count': pause_count,
            'assessment': assessment,
            'timestamps': pause_timestamps
        }
    
    except Exception as e:
        st.error(f"Error detecting pauses: {str(e)}")
        return {
            'count': 0,
            'assessment': "Not analyzed",
            'timestamps': []
        }

def calculate_wpm_with_gemini(audio_file, transcription, duration):
    """Calculate words per minute and assess ideal word count using Gemini 1.5 Flash"""
    temp_audio_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
            temp_audio.write(audio_file.read())
            temp_audio_path = temp_audio.name
        
        uploaded = genai.upload_file(temp_audio_path)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = f"""
        You are an audio analysis expert. Analyze the provided audio file and its transcription to calculate the speaking rate in words per minute (WPM). Follow these steps:
        1. Count the total number of words in the transcription.
        2. Determine the total duration of the audio in seconds.
        3. Calculate WPM as (total words / duration in minutes), where duration in minutes = duration in seconds / 60.
        
        Transcription:
        {transcription}
        
        Provide the response in the following format:
        Total words: [count]
        Audio duration: [seconds] seconds
        Speaking rate: [wpm] words per minute
        """
        
        response = model.generate_content([prompt, uploaded])
        result = response.text.strip()
        # print("result",result)
        total_words = 0
        audio_duration = 0.0
        wpm = 0.0
        
        if "Total words:" in result:
            words_match = re.search(r'Total words: (\d+)', result)
            if words_match:
                total_words = int(words_match.group(1))
        
        if "Audio duration:" in result:
            duration_match = re.search(r'Audio duration: ([\d.]+) seconds', result)
            if duration_match:
                audio_duration = float(duration_match.group(1))
        
        if "Speaking rate:" in result:
           
            wpm_match = re.search(r'Speaking rate: ([\d.]+) words per minute', result)
            # print("wpm ",wpm_match)
            if wpm_match:
                wpm = float(wpm_match.group(1))
        
        if total_words == 0 or audio_duration == 0.0 or wpm == 0.0:
            total_words = len(transcription.split())
            audio = AudioSegment.from_file(temp_audio_path)
            audio_duration = len(audio) / 1000.0
            wpm = (total_words / (audio_duration / 60.0)) if audio_duration > 0 else 0.0
        
        ideal_word_count = ""
        assessment = ""
        suggestion = ""
        
        if duration == "Less than 1 minute":
            ideal_min, ideal_max = 70, 140
            if total_words < ideal_min:
                assessment = "Too less"
                suggestion = f"The current word count is {total_words} words. For a <1 minute speech, ideal count is {ideal_min}-{ideal_max} words - please increase."
            elif total_words > ideal_max:
                assessment = "Too much"
                suggestion = f"The current word count is {total_words} words. For a <1 minute speech, ideal count is {ideal_min}-{ideal_max} words - please decrease."
            else:
                assessment = "Appropriate"
                suggestion = f"The current word count is {total_words} words, which is appropriate for a <1 minute speech."
        
        return {
            'wpm': round(wpm, 1),
            'total_words': total_words,
            'audio_duration': audio_duration,
            'word_count_assessment': assessment,
            'word_count_suggestion': suggestion
        }
    
    except Exception as e:
        st.error(f"Error calculating WPM: {str(e)}")
        try:
            total_words = len(transcription.split())
            audio = AudioSegment.from_file(temp_audio_path)
            audio_duration = len(audio) / 1000.0
            wpm = (total_words / (audio_duration / 60.0)) if audio_duration > 0 else 0.0
            return {
                'wpm': round(wpm, 1),
                'total_words': total_words,
                'audio_duration': audio_duration,
                'word_count_assessment': 'Not analyzed',
                'word_count_suggestion': 'Error in analysis.'
            }
        except Exception as inner_e:
            st.error(f"Fallback WPM calculation failed: {str(inner_e)}")
            return {
                'wpm': 0.0,
                'total_words': 0,
                'audio_duration': 0.0,
                'word_count_assessment': 'Not analyzed',
                'word_count_suggestion': 'Error in analysis.'
            }
    
    finally:
        if temp_audio_path and os.path.exists(temp_audio_path):
            try:
                os.remove(temp_audio_path)
            except Exception as e:
                st.warning(f"Failed to delete temporary file {temp_audio_path}: {str(e)}")

def analyze_speech_with_gemini(transcription, topic=None, duration=None, audio_file=None, gender=None):
    """Analyze speech using Gemini, with pause detection and WPM calculation"""
    try:
        filler_word_list = [
            "um", "uh", "like", "you know", "so", "actually", "basically", 
            "seriously", "literally", "well", "okay", "right", "I mean", 
            "sort of", "stuff", "things", "anyway", "hmm", "ah", 
            "er", "mm-hmm", "uh-huh", "just", "maybe", "probably", "guess"
        ]

        analysis_prompt = f"""You are a professional speech analyst. Your task is to analyze the following speech and provide specific, numerical feedback, including the location of filler words in the input and suggestions for limiting filler words.

Transcription:
{transcription}

Topic: {topic or 'Not specified'}
Gender: {gender or 'Not specified'}

IMPORTANT: You MUST provide actual numerical values and specific feedback. Do NOT use placeholders or empty values.

Format your analysis as follows:

Pronunciation
Overall speech accuracy: [percentage]%
Feedback: [detailed feedback on pronunciation clarity and issues]

Mood
Primary emotion: [emotion, e.g., Confidence]
Formality level: [e.g., Professional]
Audience suitability: [e.g., Highly suitable for academic audience]
Assessment: [evaluation of emotional tone and suitability]
Reasons:
- [Reason 1, e.g., Consistent use of professional terminology]
- [Reason 2, e.g., Measured pace shows confidence]
- [Reason 3, e.g., Clear emphasis on key points]

Speaking Style
Filler Words Analysis:
- [filler word, e.g., "um"]: [number] occurrences
- [filler word, e.g., "like"]: [number] occurrences
Total filler words: [total count]
Filler word percentage: [percentage]%
Location:
  [e.g., Beginning: "Hi, um, my name is Anushka Tandan"]
  [e.g., After "Tandan": "um and I would like"]
Assessment: [evaluation of filler word usage]
Suggestions for Limiting Filler Words:
1. [Suggestion 1, e.g., Pause Intentionally: Instead of saying "um," simply pause.]
2. [Suggestion 2, e.g., Plan Your Opening]
3. [Suggestion 3, e.g., Record Yourself and Listen Back]
4. [Suggestion 4, e.g., Practice Concise Speaking]

Pitch
Pitch Analysis:
Pitch variation: [e.g., Good variation, Excessive variation, or Too monotonous]
Consistency: [e.g., Very consistent with appropriate emphasis]
Average pitch: [Qualitative assessment, e.g., Too high for the given gender, Too low for the given gender, Appropriate for the given gender]

Example:
Pronunciation
Overall speech accuracy: 85%
Feedback: Clear enunciation with good articulation. Some difficulty with complex words like "statistics" and "phenomenon".

Mood
Primary emotion: Confidence
Formality level: Professional
Audience suitability: Highly suitable for academic audience
Assessment: The emotional tone effectively conveys authority and expertise
Reasons:
- Consistent use of professional terminology
- Measured pace shows confidence
- Clear emphasis on key points

Speaking Style
Filler Words Analysis:
- "um": 3 occurrences
- "like": 2 occurrences
Total filler words: 2
Filler word percentage: 2.1%
Location:
  Beginning: "Hi, um, my name is Anushka Tandan"
  After "Tandan": "um and I would like"
Assessment: Moderate use of filler words disrupts flow
Suggestions for Limiting Filler Words:
1. Pause Intentionally: Instead of saying "um," simply pause.
2. Plan Your Opening: A well-rehearsed opening reduces filler words.
3. Record Yourself: Listen to identify filler word patterns.
4. Practice Concise Speaking: Be direct in communication.

Pitch
Pitch Analysis:
Pitch variation: Good variation
Consistency: Very consistent with appropriate emphasis
Average pitch: Appropriate for the given gender

By implementing these strategies, you can significantly reduce your use of filler words and communicate with more clarity and confidence. Do not copy the examples - provide your own analysis based on the actual speech content.
"""

        model = genai.GenerativeModel('gemini-1.5-flash')
        analysis_response = model.generate_content(analysis_prompt)
        analysis_text = analysis_response.text
        
        accuracy_match = re.search(r'Overall speech accuracy: (\d+)%', analysis_text)
        accuracy = int(accuracy_match.group(1)) if accuracy_match else 0
        
        pitch_variation_match = re.search(r'Pitch variation: (.+?)\n', analysis_text)
        pitch_variation = pitch_variation_match.group(1).strip() if pitch_variation_match else 'Not analyzed'
        
        consistency_match = re.search(r'Consistency: (.+?)\n', analysis_text)
        consistency = consistency_match.group(1).strip() if consistency_match else 'Not analyzed'
        
        average_pitch_match = re.search(r'Average pitch: (.+?)(?:\n|$)', analysis_text)
        average_pitch = average_pitch_match.group(1).strip() if average_pitch_match else 'Not analyzed'

        filler_words = {word: 0 for word in filler_word_list}
        transcription_lower = transcription.lower()
        for word in filler_word_list:
            pattern = r'\b' + re.escape(word) + r'\b'
            count = len(re.findall(pattern, transcription_lower))
            filler_words[word] = count
        
        filler_words = {k: v for k, v in filler_words.items() if v > 0}
        total_fillers = sum(filler_words.values())
        total_words = len(transcription.split())
        filler_percentage = (total_fillers / total_words * 100) if total_words > 0 else 0
        
        filler_locations = []
        filler_suggestions = []
        filler_assessment = ''
        if 'Filler Words Analysis:' in analysis_text:
            filler_section = analysis_text.split('Filler Words Analysis:')[1].split('Pitch')[0]
            location_section = filler_section.split('Location:')[1].split('Assessment:')[0].strip() if 'Location:' in filler_section else ''
            for line in location_section.split('\n'):
                line = line.strip()
                if line:
                    filler_locations.append(line)
            
            filler_assessment = filler_section.split('Assessment:')[1].split('Suggestions')[0].strip() if 'Assessment:' in filler_section else ''
            
            suggestions_section = filler_section.split('Suggestions for Limiting Filler Words:')[1].strip() if 'Suggestions for Limiting Filler Words:' in filler_section else ''
            for line in suggestions_section.split('\n'):
                line = line.strip()
                if line and not line.startswith('By implementing these strategies'):
                    filler_suggestions.append(line)
        
        mood_section = analysis_text.split('Mood')[1].split('Speaking Style')[0] if 'Mood' in analysis_text and 'Speaking Style' in analysis_text else ''
        primary_emotion = mood_section.split('Primary emotion:')[1].split('\n')[0].strip() if 'Primary emotion:' in mood_section else ''
        formality = mood_section.split('Formality level:')[1].split('\n')[0].strip() if 'Formality level:' in mood_section else ''
        audience_suitability = mood_section.split('Audience suitability:')[1].split('\n')[0].strip() if 'Audience suitability:' in mood_section else ''
        mood_assessment = mood_section.split('Assessment:')[1].split('Reasons:')[0].strip() if 'Assessment:' in mood_section and 'Reasons:' in mood_section else ''
        
        reasons = []
        if 'Reasons:' in mood_section:
            reasons_section = mood_section.split('Reasons:')[1].strip()
            for line in reasons_section.split('\n'):
                line = line.strip()
                if line.startswith('-') and line.strip('-').strip():
                    reasons.append(line.strip('-').strip())
        
        wpm_data = {'wpm': 0.0, 'total_words': 0, 'audio_duration': 0.0, 'word_count_assessment': 'Not analyzed', 'word_count_suggestion': 'Not analyzed'}
        if audio_file and duration:
            audio_file.seek(0)
            wpm_data = calculate_wpm_with_gemini(audio_file, transcription, duration)
        
        pause_data = {'count': 0, 'assessment': 'Not analyzed', 'timestamps': []}
        if audio_file and duration:
            audio_file.seek(0)
            pause_data = detect_pauses_with_gemini(audio_file, duration)
        
        analysis_dict = {
            'pronunciation': {
                'accuracy': accuracy,
                'feedback': analysis_text.split('Feedback:')[1].split('\n\n')[0].strip() if 'Feedback:' in analysis_text else '',
                'difficult_words': {}
            },
            'pitch': {
                'variation': pitch_variation,
                'consistency': consistency,
                'average': average_pitch
            },
            'speech_rate': {
                'wpm': wpm_data['wpm'],
                'total_words': wpm_data['total_words'],
                'word_count_assessment': wpm_data['word_count_assessment'],
                'word_count_suggestion': wpm_data['word_count_suggestion'],
                'pauses': {
                    'count': pause_data['count'],
                    'assessment': pause_data['assessment']
                },
                'filler_words': filler_words,
                'filler_word_percentage': filler_percentage,
                'total_filler_words': total_fillers,
                'filler_locations': filler_locations,
                'filler_assessment': filler_assessment,
                'filler_suggestions': filler_suggestions
            },
            'mood': {
                'primary_emotion': primary_emotion,
                'formality': formality,
                'audience_suitability': audience_suitability,
                'mood_suitability_assessment': {
                    'assessment': mood_assessment,
                    'reasons': reasons
                }
            }
        }
        
        analysis_dict['raw_analysis'] = analysis_text
        return analysis_dict
            
    except Exception as e:
        st.error(f"Error in speech analysis: {str(e)}")
        return None

def process_with_gemini(transcription, purpose, audience, duration, tone, additional_requirements, topic, speech_analysis):
    """Generate detailed feedback using Gemini"""
    try:
        if not transcription or not transcription.strip():
            return "No transcription provided", "", "Please provide a valid transcription."

        transcription = transcription.strip()
        mispronounced_words = []
        if speech_analysis and isinstance(speech_analysis, dict):
            pronunciation = speech_analysis.get('pronunciation', {})
            if isinstance(pronunciation, dict):
                mispronounced_words = pronunciation.get('difficult_words', [])
        
        system_prompt = """You are a professional speech coach and analyst. 
        Your task is to provide detailed feedback on the speech delivery and content.
        Focus on making the feedback comprehensive and actionable.
        
        Format your feedback in the following sections:
        1. Content Analysis
        2. Delivery Analysis
        3. Specific Recommendations
        4. Overall Assessment
        """
        
        user_prompt = f"""
        Please provide detailed feedback on this speech:

        Original Text: {transcription}
        Topic: {topic or 'Not specified'}
        Purpose: {purpose or 'Not specified'}
        Target Audience: {audience or 'Not specified'}
        Duration: {duration or 'Not specified'}
        Tone: {tone or 'Not specified'}
        Additional Requirements: {additional_requirements or 'None'}
        Mispronounced Words: {mispronounced_words}

        Provide comprehensive feedback covering:
        1. Content Analysis:
           - Message clarity
           - Structure and organization
           - Relevance to topic and audience
           - Use of supporting evidence/examples
        
        2. Delivery Analysis:
           - Speaking pace and rhythm
           - Voice modulation and emphasis
           - Pronunciation and articulation
           - Use of pauses and transitions
        
        3. Specific Recommendations:
           - Areas for improvement
           - Specific techniques to try
           - Content adjustments
           - Delivery enhancements
        
        4. Overall Assessment:
           - Strengths
           - Areas for growth
           - Next steps for improvement
        """
        
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content([system_prompt, user_prompt])
        feedback = response.text
        return transcription, transcription, feedback

    except Exception as e:
        return transcription, transcription, f"Error in generating feedback: {str(e)}"



def identify_mispronounced_words(audio_file, transcription):
    """Identify mispronounced words from an audio file using Gemini and return feedback."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
            temp_audio.write(audio_file.read())
            temp_audio_path = temp_audio.name
        
        uploaded = genai.upload_file(temp_audio_path)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = """
        You are a pronunciation expert analyzing a speech audio file and its transcription. Your task is to:
        1. Evaluate the pronunciation accuracy of the speaker based on standard American English.
        2. Identify all mispronounced words, providing details for each, but only if the errors are clear and significant.

        Transcription:
        {transcription}

        For each mispronounced word, provide:
        - Word: The mispronounced word.
        - Correct Pronunciation: A clear, syllable-by-syllable breakdown (e.g., YOU-ni-VER-si-tee).
        - Phonetic Transcription: The IPA transcription (e.g., /ˌjuːnɪˈvɜːrsəti/).
        - Explanation: Why the word was mispronounced and how to correct it.

        Format your response as follows:
        Pronunciation Feedback: [Overall assessment of pronunciation accuracy]
        Mispronounced Words:
        - Word: [word]
          Correct Pronunciation: [syllable breakdown]
          Phonetic Transcription: [IPA]
          Explanation: [details]
        [If no mispronounced words are detected, state: "No mispronounced words detected."]
        Additional Notes: [Any additional observations or limitations, e.g., audio quality, need for visual confirmation, or acceptable regional variations]

        Guidelines:
        - Only identify words as mispronounced if there are clear and significant errors that impact clarity (e.g., incorrect stress, slurred syllables, wrong vowel sounds, or substitutions that change the word's meaning).
        - Do not flag minor variations, such as slight accent differences, acceptable regional pronunciations, fast speech rate, or subtle intonation differences, unless they significantly affect intelligibility.
        - If the pronunciation is clear and accurate with no significant errors, return "No mispronounced words detected" without listing any words. This is the expected outcome for accurate speech.
        - Use standard American English pronunciation as the reference unless otherwise specified.
        - Consider potential limitations, such as audio quality, background noise, or speech rate, and note these in the Additional Notes section. Do not flag issues caused by these limitations as mispronunciations unless they cause significant errors.
        - Avoid overanalyzing or assuming errors without strong evidence from the audio or transcription.
        """
        
        prompt = prompt.format(transcription=transcription)
        response = model.generate_content([prompt, uploaded])
        result = response.text.strip()
        # print("result for mis words", result)
        
        os.remove(temp_audio_path)
        
        pronunciation_feedback = ""
        mispronounced_words = []
        
        # Extract feedback
        feedback_match = re.search(r'Pronunciation Feedback: (.+?)(?:\nMispronounced Words:|$)', result, re.DOTALL)
        # print("feedbackmatch", feedback_match)
        if feedback_match:
            pronunciation_feedback = feedback_match.group(1).strip()
            # print("matched pronounced words", pronunciation_feedback)
        
        # Check for "No mispronounced words detected"
        if "No mispronounced words detected" in result:
            mispronounced_words = []
        else:
            # Extract mispronounced words using a more robust regex
            word_matches = re.findall(
                r'- Word: (.+?)\n\s*Correct Pronunciation: (.+?)\n\s*Phonetic Transcription: (.+?)\n\s*Explanation: (.+?)(?=\n\s*- Word:|\n\s*Additional Notes:|\Z)',
                result, re.DOTALL
            )
            for word, correct_pron, phonetic, explanation in word_matches:
                mispronounced_words.append({
                    'word': word.strip(),
                    'correct_pronunciation': correct_pron.strip(),
                    'phonetic_transcription': phonetic.strip(),
                    'explanation': explanation.strip()
                })
        
        # Extract additional notes
        notes_match = re.search(r'Additional Notes: (.+?)(?:\n\n|$)', result, re.DOTALL)
        notes = notes_match.group(1).strip() if notes_match else ""
        
        print("pronouced words", {
            'mispronounced_words': mispronounced_words,
            'pronunciation_feedback': pronunciation_feedback,
            'notes': notes
        })
        
        return {
            'mispronounced_words': mispronounced_words,
            'pronunciation_feedback': pronunciation_feedback,
            'notes': notes
        }
    
    except Exception as e:
        st.error(f"Error identifying mispronounced words: {str(e)}")
        return {
            'mispronounced_words': [],
            'pronunciation_feedback': 'Analysis failed due to an error.',
            'notes': str(e)
        }



# def identify_mispronounced_words(audio_file, transcription):
#     """Identify mispronounced words from an audio file using Gemini and return feedback."""
#     try:
#         with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
#             temp_audio.write(audio_file.read())
#             temp_audio_path = temp_audio.name
        
#         uploaded = genai.upload_file(temp_audio_path)
#         model = genai.GenerativeModel('gemini-1.5-flash')
        
#         prompt = """
#         You are a pronunciation expert analyzing a speech audio file and its transcription. Your task is to:
#         1. Evaluate the pronunciation accuracy of the speaker based on standard American English.
#         2. Identify all mispronounced words, providing details for each, but only if the errors are clear and significant.

#         Transcription:
#         {transcription}

#         For each mispronounced word, provide:
#         - Word: The mispronounced word.
#         - Correct Pronunciation: A clear, syllable-by-syllable breakdown (e.g., YOU-ni-VER-si-tee).
#         - Phonetic Transcription: The IPA transcription (e.g., /ˌjuːnɪˈvɜːrsəti/).
#         - Explanation: Why the word was mispronounced and how to correct it.

#         Format your response as follows:
#         Pronunciation Feedback: [Overall assessment of pronunciation accuracy]
#         Mispronounced Words:
#         - Word: [word]
#           Correct Pronunciation: [syllable breakdown]
#           Phonetic Transcription: [IPA]
#           Explanation: [details]
#         [If no mispronounced words are detected, state: "No mispronounced words detected."]
#         Additional Notes: [Any additional observations or limitations, e.g., audio quality, need for visual confirmation, or acceptable regional variations]

#         Guidelines:
#         - Only identify words as mispronounced if there are clear and significant errors that impact clarity (e.g., incorrect stress, slurred syllables, wrong vowel sounds, or substitutions that change the word's meaning).
#         - Do not flag minor variations, such as slight accent differences, acceptable regional pronunciations, fast speech rate, or subtle intonation differences, unless they significantly affect intelligibility.
#         - If the pronunciation is clear and accurate with no significant errors, return "No mispronounced words detected" without listing any words. This is the expected outcome for accurate speech.
#         - Use standard American English pronunciation as the reference unless otherwise specified.
#         - Consider potential limitations, such as audio quality, background noise, or speech rate, and note these in the Additional Notes section. Do not flag issues caused by these limitations as mispronunciations unless they cause significant errors.
#         - Avoid overanalyzing or assuming errors without strong evidence from the audio or transcription.
#         """
        
#         prompt = prompt.format(transcription=transcription)
#         response = model.generate_content([prompt, uploaded])
#         result = response.text.strip()
#         print("result for mis words",result)
        
#         os.remove(temp_audio_path)
        
#         # Debug: Log transcription and raw response
#         # st.write("Transcription:", transcription)
#         # st.write("Raw Gemini Pronunciation Response:", result)
        
#         pronunciation_feedback = ""
#         mispronounced_words = []
        
#         # Extract feedback
#         feedback_match = re.search(r'Pronunciation Feedback: (.+?)(?:\nMispronounced Words:|$)', result, re.DOTALL)
#         print("feedbackmatch",feedback_match)
#         if feedback_match:
#             pronunciation_feedback = feedback_match.group(1).strip()
#             print("matched pronouced words",pronunciation_feedback)
        
#         # Check for "No mispronounced words detected"
#         if "No mispronounced words detected" in result:
#             mispronounced_words = []
#         else:
#             # Extract mispronounced words section
#             mispronounced_section = re.search(r'Mispronounced Words:\n(.+?)(?:\nAdditional Notes:|\n\n|$)', result, re.DOTALL)
#             if mispronounced_section:
#                 section_text = mispronounced_section.group(1).strip()
                
#                 # Split by word entries
#                 word_blocks = re.split(r'\n\s*- Word: ', section_text)[1:]  # Skip first empty split
#                 failed_blocks = []
#                 for block in word_blocks:
#                     block = block.strip()
#                     word_match = re.match(
#                         r'(.+?)\n\s*Correct Pronunciation: (.+?)\n\s*Phonetic Transcription: (.+?)\n\s*Explanation: (.+?)(?=\n\s*- Word:|\n\s*Additional Notes:|\n\n|$)',
#                         block, re.DOTALL
#                     )
#                     if word_match:
#                         word, correct_pron, phonetic, explanation = word_match.groups()
#                         mispronounced_words.append({
#                             'word': word.strip(),
#                             'correct_pronunciation': correct_pron.strip(),
#                             'phonetic_transcription': phonetic.strip(),
#                             'explanation': explanation.strip()
#                         })
#                     else:
#                         failed_blocks.append(block)
                
#                 # Debug: Log failed blocks
#                 if failed_blocks:
#                     st.write("Failed to parse blocks:", failed_blocks)
                
#                 # Fallback parsing if any entries were missed
#                 if len(mispronounced_words) < 13:  # Expecting 13 words
#                     word_matches = re.findall(
#                         r'- Word: (.+?)\n\s*Correct Pronunciation: (.+?)\n\s*Phonetic Transcription: (.+?)\n\s*Explanation: (.+?)(?=\n\s*- Word:|\n\s*Additional Notes:|\n\n|$)',
#                         section_text, re.DOTALL
#                     )
#                     for word, correct_pron, phonetic, explanation in word_matches:
#                         entry = {
#                             'word': word.strip(),
#                             'correct_pronunciation': correct_pron.strip(),
#                             'phonetic_transcription': phonetic.strip(),
#                             'explanation': explanation.strip()
#                         }
#                         if entry not in mispronounced_words:
#                             mispronounced_words.append(entry)
        
#         # Debug: Log parsed mispronounced words
#         # st.write("Parsed Mispronounced Words:", mispronounced_words)
        
#         # Extract additional notes
#         notes_match = re.search(r'Additional Notes: (.+?)(?:\n\n|$)', result, re.DOTALL)
#         notes = notes_match.group(1).strip() if notes_match else ""
        
#         return {
#             'mispronounced_words': mispronounced_words,
#             'pronunciation_feedback': pronunciation_feedback,
#             'notes': notes
#         }
    
#     except Exception as e:
#         st.error(f"Error identifying mispronounced words: {str(e)}")
#         return {
#             'mispronounced_words': [],
#             'pronunciation_feedback': 'Analysis failed due to an error.',
#             'notes': str(e)
#         }



def process_with_gpt(openai_api_key, transcription, purpose, audience, duration, tone, additional_requirements, topic, speech_analysis):
    """Process transcription and generate refined speech using GPT, synchronizing bold markers."""
    try:
        if not transcription or not transcription.strip():
            return "No transcription provided", "", "Please provide a valid transcription."

        transcription = transcription.strip()
        mispronounced_words = []
        if speech_analysis and isinstance(speech_analysis, dict):
            pronunciation = speech_analysis.get('pronunciation', {})
            if isinstance(pronunciation, dict):
                mispronounced_words = pronunciation.get('difficult_words', [])

        system_prompt = """You are a professional speech writer and coach. 
        Your task is to refine the given speech and provide detailed feedback on its delivery.
        Focus on making the speech more engaging and appropriate for the specified audience and purpose.
        
        For the transcription and refined speech:
        1. Add **bold** around words that should be emphasized
        2. Add | (vertical bar) where the speaker should pause
        3. Mark mispronounced words with <mispronounced> tags
        
        Format your response exactly as follows:
        ORIGINAL:
        [transcription with markers]
        
        REFINED:
        [refined speech with markers]
        
        FEEDBACK:
        [detailed feedback]
        """

        user_prompt = f"""
        Please analyze and refine this speech:

        Original Text: {transcription}
        Topic: {topic or 'Not specified'}
        Purpose: {purpose or 'Not specified'}
        Target Audience: {audience or 'Not specified'}
        Duration: {duration or 'Not specified'}
        Tone: {tone or 'Not specified'}
        Mispronounced Words: {mispronounced_words}

        Provide:
        1. The original transcription with emphasis markers and pause indicators
        2. A refined version of the speech with the same markers
        3. Detailed feedback on delivery and content
        """

        groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=1000,
        )

        full_response = completion.choices[0].message.content
        parts = full_response.split('\n\n')
        original = ""
        refined = ""
        feedback = ""

        for part in parts:
            if part.startswith('ORIGINAL:'):
                original = part.replace('ORIGINAL:', '').strip()
            elif part.startswith('REFINED:'):
                refined = part.replace('REFINED:', '').strip()
            elif part.startswith('FEEDBACK:'):
                feedback = part.replace('FEEDBACK:', '').strip()

        if not original or not refined:
            original = transcription
            words = original.split()
            for i, word in enumerate(words):
                if len(word) > 4 and word[0].isupper():
                    words[i] = f"**{word}**"
            original = ' '.join(words)
            return original, transcription, "Could not process the response properly. Using original transcription with basic emphasis."

        bolded_words = []
        bold_pattern = r'\*\*(.*?)\*\*'
        for match in re.finditer(bold_pattern, refined):
            bolded_word = match.group(1).strip()
            bolded_words.append(bolded_word.lower())

        def apply_bold_to_word(word):
            if word.lower() in bolded_words:
                return f"**{word}**"
            return word

        parts = re.split(r'(\s+|\|)', original)
        modified_parts = []
        for part in parts:
            if part.strip() and part not in ['|', ' ']:
                modified_parts.append(apply_bold_to_word(part))
            else:
                modified_parts.append(part)
        
        modified_original = ''.join(modified_parts)

        return modified_original, refined, feedback

    except Exception as e:
        return transcription, transcription, f"Error in processing: {str(e)}"

def generate_audio_from_text(client, text, filename="temp_output.wav"):
    try:
        if not text or not text.strip():
            st.error("No text provided for audio generation")
            return None

        speak_options = {'text': text.strip()}
        options = SpeakOptions(model="aura-asteria-en", encoding="linear16", container="wav")
        response = client.speak.v("1").save(filename, speak_options, options)
        
        with open(filename, 'rb') as audio_file:
            audio_bytes = audio_file.read()
            audio_base64 = base64.b64encode(audio_bytes).decode('utf-8')
        
        if os.path.exists(filename):
            os.remove(filename)
            
        return audio_base64
    except Exception as e:
        st.error(f"Error generating audio: {e}")
        return None

def save_processed_data(session_id, data_type, content):
    """Save processed data to local storage"""
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    if data_type == 'text':
        filepath = f'processed_data/text/{session_id}_{timestamp}.txt'
        with open(filepath, 'w') as f:
            f.write(content)
    elif data_type == 'audio':
        filepath = f'processed_data/audio/{session_id}_{timestamp}.wav'
        with open(filepath, 'wb') as f:
            f.write(content)
    return filepath

def load_processed_data(filepath):
    """Load processed data from local storage"""
    if not os.path.exists(filepath):
        return None
    
    if filepath.endswith('.txt'):
        with open(filepath, 'r') as f:
            return f.read()
    elif filepath.endswith('.wav'):
        with open(filepath, 'rb') as f:
            return f.read()
    return None

def convert_mp3_to_wav(mp3_file):
    """Convert MP3 file to WAV format"""
    try:
        audio = AudioSegment.from_mp3(mp3_file)
        wav_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
        audio.export(wav_file.name, format="wav")
        return wav_file.name
    except Exception as e:
        st.error(f"Error converting MP3 to WAV: {str(e)}")
        return None

def convert_mp4_to_wav(mp4_file):
    """Convert MP4 file to WAV format"""
    try:
        if not mp4_file:
            st.error("No file provided for conversion")
            return None

        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp_mp4:
            tmp_mp4.write(mp4_file.getvalue())
            mp4_path = tmp_mp4.name

        wav_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
        wav_path = wav_file.name
        audio = AudioSegment.from_file(mp4_path, format="mp4")
        audio = audio.set_frame_rate(44100).set_channels(1)
        audio.export(wav_path, format="wav")
        
        os.unlink(mp4_path)
        return wav_path
            
    except Exception as e:
        if 'mp4_path' in locals() and os.path.exists(mp4_path):
            os.unlink(mp4_path)
        if 'wav_path' in locals() and os.path.exists(wav_path):
            os.unlink(wav_path)
        st.error(f"Error converting MP4 to WAV: {str(e)}")
        return None

def generate_word_pronunciation(word):
    """Generate audio for a single word using gTTS"""
    try:
        # tts = gTTS(text=word, lang='en')
        tts = tts = gTTS(text = word, lang='en', tld='co.in')
        audio_io = io.BytesIO()
        tts.write_to_fp(audio_io)
        audio_io.seek(0)
        return audio_io
    except Exception as e:
        st.error(f"Error generating pronunciation for {word}: {str(e)}")
        return None

# def format_transcription_text(text, mispronounced_words=None):
#     """Format transcription text with emphasis markers and mispronounced words"""
#     if mispronounced_words is None:
#         mispronounced_words = []
    
#     # Convert bold markers to HTML
#     pattern = r'\*\*(.*?)\*\*'
#     text = re.sub(pattern, r'<span class="bold-word">\1</span>', text)
#     # Convert pause markers to HTML
#     text = text.replace('|', '<span class="pause-marker">|</span>')
    
#     # Split text, preserving HTML tags and spaces
#     parts = re.split(r'(\s+|<[^>]*>)', text)
#     formatted_parts = []
    
#     for part in parts:
#         if part.strip() and not part.startswith('<'):
#             word = part.strip()
#             # Remove punctuation for mispronounced word comparison
#             clean_word = re.sub(r'[^\w\s]', '', word).lower()
#             if any(clean_word == mw.lower() for mw in mispronounced_words):
#                 formatted_parts.append(f'<span class="mispronounced">{word}</span>')
#             else:
#                 formatted_parts.append(word)
#         else:
#             formatted_parts.append(part)
    
#     return ''.join(formatted_parts)

def format_transcription_text(text, mispronounced_words=None):
    if mispronounced_words is None:
        mispronounced_words = []
    
    # Normalize mispronounced words
    mispronounced_words_normalized = [re.sub(r'[^\w\s]', '', word).lower() for word in mispronounced_words]
    
    # Debug: Log normalized mispronounced words
    # st.write("Normalized Mispronounced Words (Text):", mispronounced_words_normalized)
    
    # Convert bold markers to HTML
    pattern = r'\*\*(.*?)\*\*'
    text = re.sub(pattern, r'<span class="bold-word">\1</span>', text)
    # Convert pause markers to HTML
    text = text.replace('|', '<span class="pause-marker">|</span>')
    
    # Split text, preserving HTML tags, spaces, and punctuation
    parts = re.split(r'(\s+|<[^>]*>|[.,!?;])', text)
    formatted_parts = []
    
    for part in parts:
        if part.strip() and not part.startswith('<') and part not in ['.', ',', '!', '?', ';']:
            word = part.strip()
            word_normalized = re.sub(r'[^\w\s]', '', word).lower()
            if word_normalized in mispronounced_words_normalized:
                formatted_parts.append(f'<span class="mispronounced">{word}</span>')
            else:
                formatted_parts.append(word)
        else:
            formatted_parts.append(part)
    
    return ''.join(formatted_parts)


# def format_transcription_with_emphasis(transcription, mispronounced_words=None):
#     """Format transcription with emphasis markers and mispronounced words"""
#     if mispronounced_words is None:
#         mispronounced_words = []
    
#     pattern = r'\*\*(.*?)\*\*'
#     transcription = re.sub(pattern, r'<span class="bold-word">\1</span>', transcription)
    
#     parts = re.split(r'(\s+|<[^>]*>)', transcription)
#     formatted_parts = []
    
#     for part in parts:
#         if part.strip() and not part.startswith('<'):
#             word = part.strip()
#             word_lower = word.lower()
#             if word_lower in [mw.lower() for mw in mispronounced_words]:
#                 formatted_parts.append(f'<span class="mispronounced">{word}</span>')
#             else:
#                 formatted_parts.append(word)
#         else:
#             formatted_parts.append(part)
    
#     formatted_text = ''.join(formatted_parts)
    
#     legend = """
#     <div class="transcription-legend">
#         <strong>Legend:</strong><br>
#         <span class="bold-word">Bold words</span> - Words to emphasize<br>
#         <span class="pause-marker">|</span> - Pause in speech<br>
#         <span class="mispronounced">Highlighted words</span> - Words that need pronunciation improvement
#     </div>
#     """
    
#     return f"""
#     <div class="transcription-container">
#         {legend}
#         <div class="transcription-text">
#             {formatted_text}
#         </div>
#     </div>
#     """


def format_transcription_with_emphasis(transcription, mispronounced_words=None):
    if mispronounced_words is None:
        mispronounced_words = []
    
    # Normalize mispronounced words (lowercase, remove punctuation)
    mispronounced_words_normalized = [re.sub(r'[^\w\s]', '', word).lower() for word in mispronounced_words]
    
    # Debug: Log normalized mispronounced words
    # st.write("Normalized Mispronounced Words:", mispronounced_words_normalized)
    
    # Convert bold markers to HTML
    pattern = r'\*\*(.*?)\*\*'
    transcription = re.sub(pattern, r'<span class="bold-word">\1</span>', transcription)
    # Convert pause markers to HTML
    transcription = transcription.replace('|', '<span class="pause-marker">|</span>')
    
    # Split text, preserving HTML tags, spaces, and punctuation
    parts = re.split(r'(\s+|<[^>]*>|[.,!?;])', transcription)
    formatted_parts = []
    
    for part in parts:
        if part.strip() and not part.startswith('<') and part not in ['.', ',', '!', '?', ';']:
            word = part.strip()
            # Normalize word for comparison (remove punctuation, lowercase)
            word_normalized = re.sub(r'[^\w\s]', '', word).lower()
            if word_normalized in mispronounced_words_normalized:
                formatted_parts.append(f'<span class="mispronounced">{word}</span>')
            else:
                formatted_parts.append(word)
        else:
            formatted_parts.append(part)
    
    formatted_text = ''.join(formatted_parts)
    
    legend = """
    <div class="transcription-legend">
        <strong>Legend:</strong><br>
        <span class="bold-word">Bold words</span> - Words to emphasize<br>
        <span class="pause-marker">|</span> - Pause in speech<br>
        <span class="mispronounced">Highlighted words</span> - Words that need pronunciation improvement
    </div>
    """
    
    return f"""
    <div class="transcription-container">
        {legend}
        <div class="transcription-text">
            {formatted_text}
        </div>
    </div>
    """

def format_detailed_feedback(results, pronunciation_data=None):
    """Format detailed feedback sections with dynamic data, including pronunciation feedback."""
    if isinstance(results, str):
        return f"""
        <div class="content-section">
            <h2>Detailed Feedback</h2>
            <pre>{results}</pre>
        </div>
        """
    
    pronunciation = results.get('pronunciation', {})

    pitch = results.get('pitch', {})
    # print("pitch",pitch)
    
    speech_rate = results.get('speech_rate', {})
    # print("speech_rate",speech_rate)
    mood = results.get('mood', {})
    
    if not isinstance(pronunciation_data, dict):
        pronunciation_data = {
            'mispronounced_words': [],
            'pronunciation_feedback': 'Not analyzed (invalid pronunciation data)',
            'notes': ''
        }
        st.error("Invalid pronunciation data format. Expected a dictionary.")
    
    # Debug: Log pronunciation_data
    # st.write("Pronunciation Data for Rendering:", pronunciation_data)
    
    pronunciation_section = """
    <h3>Pronunciation</h3>
    <ul>
        <li><strong>Overall speech accuracy:</strong> {accuracy}%</li>
        <li><strong>Pronunciation Feedback:</strong> {feedback}</li>
    """.format(
        accuracy=pronunciation.get('accuracy', 0),
        feedback=pronunciation_data.get('pronunciation_feedback', 'Not analyzed')
    )
    
    mispronounced_words = pronunciation_data.get('mispronounced_words', [])
    if mispronounced_words:
        pronunciation_section += "<li><strong>Mispronounced Words:</strong></li><ul>"
        for item in mispronounced_words:
            # audio_io = generate_word_pronunciation(item['word'])
            # audio_html = ""
            # if audio_io:
            #     audio_bytes = audio_io.getvalue()
            #     audio_base64 = base64.b64encode(audio_bytes).decode('utf-8')
            #     audio_html = f"""
            #     <audio class="audio-player" controls>
            #         <source src="data:audio/mp3;base64,{audio_base64}" type="audio/mp3">
            #         Your browser does not support the audio element.
            #     </audio>
            #     """
            pronunciation_section += f"""
            <li>
                Word: {item['word']}<br>
                Correct Pronunciation: {item['correct_pronunciation']}<br>
                Phonetic Transcription: {item['phonetic_transcription']}<br>
                Explanation: {item['explanation']}<br>
               
            </li>
            """
        pronunciation_section += "</ul>"
    else:
        pronunciation_section += "<li><strong>Mispronounced Words:</strong> No mispronounced words detected.</li>"
    
    notes = pronunciation_data.get('notes', '')
    if notes:
        pronunciation_section += f"<li><strong>Additional Notes:</strong> {notes}</li>"
    
    pronunciation_section += "</ul>"
    
    mood_section = """
    <h3>Mood</h3>
    <ul>
        <li><strong>Primary emotion:</strong> {primary_emotion}</li>
        <li><strong>Formality level:</strong> {formality}</li>
        <li><strong>Audience suitability:</strong> {audience}</li>
        <li><strong>Assessment:</strong> {assessment}</li>
    """.format(
        primary_emotion=mood.get('primary_emotion', 'Not analyzed'),
        formality=mood.get('formality', 'Not analyzed'),
        audience=mood.get('audience_suitability', 'Not analyzed'),
        assessment=mood.get('mood_suitability_assessment', {}).get('assessment', 'Not analyzed')
    )
    
    reasons = mood.get('mood_suitability_assessment', {}).get('reasons', [])
    if reasons:
        mood_section += "<li><strong>Reasons:</strong></li><ul>"
        for reason in reasons:
            mood_section += f"<li>{reason}</li>"
        mood_section += "</ul>"
    else:
        mood_section += "<li><strong>Reasons:</strong> No specific reasons provided.</li>"
    
    mood_section += "</ul>"
    
    speaking_style_section = """
    <h3>Speaking Style</h3>
    <ul>
        <li><strong>Speaking rate:</strong> {rate} words per minute</li>
        <li><strong>Ideal word count:</strong> {word_count_assessment} ({word_count_suggestion})</li>
        <li><strong>Pauses:</strong> {pause_assessment}</li>
    """.format(
        rate=speech_rate.get('wpm', 0),
        word_count_assessment=speech_rate.get('word_count_assessment', 'Not analyzed'),
        word_count_suggestion=speech_rate.get('word_count_suggestion', 'Not analyzed'),
        pause_assessment=speech_rate.get('pauses', {}).get('assessment', 'Not analyzed')
    )
    
    if speech_rate.get('filler_words'):
        speaking_style_section += "<li><strong>Filler words analysis:</strong></li><ul>"
        for word, count in speech_rate['filler_words'].items():
            speaking_style_section += f"<li>'{word}': {count} occurrences</li>"
        speaking_style_section += "</ul>"
    speaking_style_section += f"<li><strong>Total filler words:</strong> {speech_rate.get('total_filler_words', 0)}</li>"
    
    if speech_rate.get('filler_locations'):
        speaking_style_section += "<li><strong>Location:</strong></li><ul>"
        for location in speech_rate['filler_locations']:
            speaking_style_section += f"<li>{location}</li>"
        speaking_style_section += "</ul>"
    
    if speech_rate.get('filler_assessment'):
        speaking_style_section += f"<li><strong>Assessment:</strong> {speech_rate['filler_assessment']}</li>"
    
    if speech_rate.get('filler_suggestions'):
        speaking_style_section += "<li><strong>Suggestions for Limiting Filler Words:</strong></li><ul>"
        for suggestion in speech_rate['filler_suggestions']:
            speaking_style_section += f"<li>{suggestion}</li>"
        speaking_style_section += "</ul>"
        speaking_style_section += "<li>By implementing these strategies, you can significantly reduce your use of filler words and communicate with more clarity and confidence.</li>"
    
    speaking_style_section += "</ul>"
    
    pitch_section = """
    <h3>Pitch</h3>
    <ul>
        <li><strong>Pitch Analysis:</strong></li>
        <ul>
            <li>Detected Gender: {detected_gender}</li>
            <li>Pitch variation: {variation}</li>
            <li>Consistency: {consistency}</li>
            <li>Average pitch: {average}</li>
        </ul>
    </ul>
    """.format(
        detected_gender=pitch.get('detected_gender', 'Not analyzed'),
        variation=pitch.get('variation', 'Not analyzed'),
        consistency=pitch.get('consistency', 'Not analyzed'),
        average=pitch.get('average', 'Not analyzed')
    )
    
    return f"""
    <div class="content-section">
        <h2>Detailed Feedback</h2>
        {pronunciation_section}
        {mood_section}
        {speaking_style_section}
        {pitch_section}
    </div>
    """

def analyze_audio_gender_and_pitch(audio_file):
    """Analyze the gender and pitch of an audio file using Gemini."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as temp_audio:
            temp_audio.write(audio_file.read())
            temp_audio_path = temp_audio.name
        
        uploaded = genai.upload_file(temp_audio_path)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = """
        You are an audio analysis expert specializing in voice characteristics. Your task is to analyze the provided audio file to:
        1. Detect the speaker's gender (male, female, or indeterminate if unclear).
        2. Evaluate the pitch characteristics, including:
           - Whether the average pitch is too high, too low, or appropriate for the detected gender.
           - Whether the pitch variation is excessive, too monotonous, or appropriate.

        Provide the response in the following format:
        Detected Gender: [male, female, or indeterminate]
        Average Pitch: [Too high for the given gender, Too low for the given gender, Appropriate for the given gender, or Cannot be definitively determined without clear gender detection]
        Pitch Variation: [Excessive, Too monotonous, or Appropriate]

        Guidelines:
        - For gender detection, use vocal characteristics such as pitch, timbre, and resonance. Typical ranges are:
          - Male: 85-180 Hz (lower pitch, deeper resonance).
          - Female: 165-255 Hz (higher pitch, lighter resonance).
          - Indeterminate: If characteristics overlap significantly or are unclear.
        - For average pitch:
          - Compare the speaker's average pitch to the typical range for the detected gender.
          - Male: Classify as "Too high" if >190 Hz, "Too low" if <80 Hz, "Appropriate" if 80-190 Hz.
          - Female: Classify as "Too high" if >260 Hz, "Too low" if <150 Hz, "Appropriate" if 150-260 Hz.
          - Indeterminate: Use a neutral range (120-220 Hz) for assessment.
        - For pitch variation:
          - "Excessive" if pitch fluctuates >80 Hz across the audio, indicating overly dramatic shifts.
          - "Too monotonous" if pitch variation is <15 Hz, indicating a flat delivery.
          - "Appropriate" if variation is 15-80 Hz, supporting natural, engaging speech.
        - If gender is indeterminate, provide a fallback pitch assessment based on the neutral range (120-220 Hz).
        - Be precise and avoid defaulting to "Appropriate" unless the pitch is clearly within the specified range for the gender.
        - If the audio is too short (<3 seconds) or unclear, note limitations in the analysis.
        """
        
        response = model.generate_content([prompt, uploaded])
        result = response.text.strip()
        
        os.remove(temp_audio_path)
        
        detected_gender = "indeterminate"
        average_pitch = "Cannot be definitively determined without clear gender detection"
        pitch_variation = "Not analyzed"
        
        gender_match = re.search(r'Detected Gender: (male|female|indeterminate)', result)
        if gender_match:
            detected_gender = gender_match.group(1)
        
        pitch_match = re.search(r'Average Pitch: (.+?)(?:\n|$)', result)
        if pitch_match:
            average_pitch = pitch_match.group(1).strip()
        
        variation_match = re.search(r'Pitch Variation: (.+?)(?:\n|$)', result)
        if variation_match:
            pitch_variation = variation_match.group(1).strip()
        
        return {
            'detected_gender': detected_gender,
            'average_pitch': average_pitch,
            'pitch_variation': pitch_variation
        }
    
    except Exception as e:
        st.error(f"Error analyzing audio gender and pitch: {str(e)}")
        return {
            'detected_gender': 'indeterminate',
            'average_pitch': 'Cannot be definitively determined due to analysis error',
            'pitch_variation': 'Not analyzed'
        }


def services():
    analyzer = SpeechAnalyzer()
    
    if 'usage_count' not in st.session_state:
        st.session_state.usage_count = 0
    
    if st.session_state.usage_count >= 5 and not st.session_state.get('is_authenticated', False):
        st.error("""
            You have reached the free trial limit of 5 uses.
            Please fill out the Contact Us form to request access credentials.
            Our team will review your request and provide you with login details.
            Thank you for trying out SpeechSmith!
        """)
        return

    load_services_css()
    
    if 'session_id' not in st.session_state:
        st.session_state.session_id = f"session_{time.strftime('%Y%m%d%H%M%S')}"
    
    if 'text_filepath' not in st.session_state:
        st.session_state.text_filepath = None
    if 'audio_filepath' not in st.session_state:
        st.session_state.audio_filepath = None
    
    st.markdown('<h1 class="gradient-text">Our Services</h1>', unsafe_allow_html=True)
    st.markdown('<p class="section-subtitle">At speechsmith, we offer a seamless and effective way to refine your speech delivery, ensuring it meets your specific goals and resonates with your audience.</p>', unsafe_allow_html=True)
    
    if not st.session_state.get('is_authenticated', False):
        remaining_uses = 5 - st.session_state.usage_count
        st.info(f"You have {remaining_uses} free trial uses remaining. Please contact us for full access.")

        content = """
        <div class="content-section"> 
        <h2>Speech Analysis and Feedback</h2>
        <p>Upload your speech as an audio or video file, and receive comprehensive feedback across several parameters:</p>
        <ul>
            <li><strong>Pronunciation:</strong> Precision scoring and suggestions to enhance clarity.</li>
            <li><strong>Posterior Score:</strong> Analysis of fluency and coherence.</li>
            <li><strong>Semantic Analysis:</strong> Checking the alignment of your speech content with your intended message and audience.</li>
            <li><strong>Words Per Minute (WPM):</strong> Insights on the ideal pace for engaging delivery.</li>
            <li><strong>Articulation Rate:</strong> Evaluating clarity and emphasis on key points.</li>
            <li><strong>Filler Words:</strong> Identifying unnecessary fillers and providing strategies to minimize them.</li>
        </ul>

        <h2>Personalized Improvement Tips</h2>
        <p>Based on the analysis, we provide actionable feedback on how you can enhance specific areas of your speech, whether that's clarity, pace, or vocabulary usage. You'll receive structured advice to help you sound more confident and professional.</p>

        <h2>Script Refinement</h2>
        <p>Receive an edited version of your script that aligns with your requirements and intended audience. Our suggestions will help tailor your content to maximize impact, ensuring that your speech resonates with your listeners.</p>

        <h2>Comprehensive Progress Reports</h2>
        <p>Track your improvement over time! Our platform keeps a record of your past uploads, allowing you to see your progress in metrics like articulation rate, pronunciation, and speech pace. With regular practice, watch your confidence grow with every upload.</p>

        <h2>Speech Crafting for Diverse Scenarios</h2>
        <p>We support users in creating and refining speeches for various purposes, including:</p>
        <ul>
            <li><strong>Debates and Competitions:</strong> Get debate-ready with speech pacing, rebuttal framing, and structured delivery feedback.</li>
            <li><strong>Presentations:</strong> Improve your presentation style for maximum engagement in team meetings, client pitches, or school projects.</li>
            <li><strong>Public Speaking Practice:</strong> For those simply wanting to refine their public speaking, SpeechSmith offers ongoing feedback to strengthen and elevate your speaking style.</li>
        </ul>
        </div>
    """
    st.html(content)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
            <div class="service-card">
                <i class="fas fa-microphone service-icon"></i>
                <h3>AUDIO UPLOAD</h3>
                <p>Easily upload your speech recordings to our platform for comprehensive analysis.</p>
            </div>
            
            <div class="service-card">
                <i class="fas fa-tasks service-icon"></i>
                <h3>CUSTOMIZED SPEECH GOALS</h3>
                <p>Tailor your speech refinement by selecting your target audience and the intent of your speech.</p>
            </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
            <div class="service-card">
                <i class="fas fa-comments service-icon"></i>
                <h3>DETAILED FEEDBACK</h3>
                <p>Receive in-depth feedback on your speech, including insights on tone, pacing, and clarity.</p>
            </div>
            
            <div class="service-card">
                <i class="fas fa-edit service-icon"></i>
                <h3>REFORMED SPEECH</h3>
                <p>Get a refined version of your speech that aligns perfectly with your chosen audience and intent.</p>
            </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    st.markdown('<h2 class="gradient-text">Upload Your Speech</h2>', unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader("Upload your audio file", type=['mp3', 'wav', 'mp4'])
    uploaded_document = st.file_uploader("Upload the script as a pdf/word document (optional, but suggested)", type=['pdf', 'docx'])
    
    if uploaded_file:
        st.subheader("Original Speech")
        st.audio(uploaded_file, format='audio/wav')
        st.session_state.original_audio = uploaded_file

    topic_of_speech = ""
    
    if uploaded_file is not None or uploaded_document is not None:
        topic_of_speech = st.text_input("Enter the topic of the speech")
        gender = st.selectbox("Select Gender", ['male', 'female', 'Other', 'Prefer not to say'])
        purpose = st.selectbox("What is the purpose of your speech?", ["Inform", "Persuade/Inspire", "Entertain", "Other"])
        if purpose == "Other":
            purpose = st.text_input("Please specify the purpose")

        audience = st.selectbox("Who is your target audience?", ["Classmates/Colleagues", "Teachers/Professors", "General public", "Other"])
        if audience == "Other":
            audience = st.text_input("Please specify the audience")

        duration = st.selectbox("How long is your speech intended to be?", ["Less than 1 minute", "1-3 minutes", "3-5 minutes", "More than 5 minutes"])
        tone = st.selectbox("What tone do you wish to adopt?", ["Formal", "Informal", "Humorous", "Other"])
        if tone == "Other":
            tone = st.text_input("Please specify the tone")
        
        additional_requirements = st.text_area("Any additional requirements or preferences? (Optional)", height=100)
        
        if st.button("Process Speech"):
            if not purpose or not audience or not duration or not tone:
                st.error("Please fill out all required fields before processing.")
                st.stop()
            
            if not st.session_state.get('is_authenticated', False):
                st.session_state.usage_count += 1
            
            status_container = st.empty()
            
            try:
                status_container.markdown("""
                    <div class="stStatusContainer stInfo">
                        <span class="step-counter">1</span>
                        Transcribing your speech using Gemini...
                    </div>
                """, unsafe_allow_html=True)
                
                if uploaded_file:
                    transcription = transcribe_audio(uploaded_file)
                    if not transcription or not transcription.strip():
                        st.error("No speech detected in the audio file")
                        st.stop()
                    
                    st.session_state.transcription = transcription
                    
                    status_container.markdown("""
                        <div class="stStatusContainer stInfo">
                            <span class="step-counter">2</span>
                            Identifying mispronounced words...
                        </div>
                    """, unsafe_allow_html=True)
                    
                    uploaded_file.seek(0)
                    pronunciation_data = identify_mispronounced_words(uploaded_file, transcription)
                    # print("pronouced words",pronunciation_data)
                    
                    status_container.markdown("""
                        <div class="stStatusContainer stInfo">
                            <span class="step-counter">3</span>
                            Analyzing gender and pitch...
                        </div>
                    """, unsafe_allow_html=True)
                    
                    uploaded_file.seek(0)
                    pitch_analysis = analyze_audio_gender_and_pitch(uploaded_file)
                    
                    try:
                        st.info("Generating AI version of your speech...")
                        tts = gTTS(text=transcription, lang='en')
                        ai_audio_io = io.BytesIO()
                        tts.write_to_fp(ai_audio_io)
                        ai_audio_io.seek(0)
                        ai_audio_bytes = ai_audio_io.getvalue()
                        
                        st.session_state.ai_audio_io = ai_audio_io
                        st.session_state.ai_audio_bytes = ai_audio_bytes
                        st.success("AI version generated successfully!")
                    except Exception as e:
                        st.error(f"Error generating AI version: {str(e)}")
                    
                    status_container.markdown("""
                        <div class="stStatusContainer stInfo">
                            <span class="step-counter">4</span>
                            Analyzing speech metrics...
                        </div>
                    """, unsafe_allow_html=True)
                    
                    results = analyze_speech_with_gemini(transcription, topic_of_speech, duration, uploaded_file, gender=pitch_analysis['detected_gender'])
                    
                    if not results:
                        st.error("Failed to analyze speech")
                        st.stop()
                    
                    # Override pitch analysis with audio-based results
                    results['pitch'] = {
                        'variation': pitch_analysis['pitch_variation'],
                        'consistency': results['pitch'].get('consistency', 'Not analyzed'),
                        'average': pitch_analysis['average_pitch'],
                        'detected_gender': pitch_analysis['detected_gender']
                    }
                    
                    # Override pronunciation feedback for consistency
                    results['pronunciation']['feedback'] = pronunciation_data['pronunciation_feedback']
                    
                else:
                    transcription = extract_text_from_document(uploaded_document)
                    if not transcription or not transcription.strip():
                        st.error("No text found in the document")
                        st.stop()
                    results = analyze_speech_with_gemini(transcription, topic_of_speech, duration, gender=gender)
                    pronunciation_data = {
                        'mispronounced_words': [],
                        'pronunciation_feedback': 'Not analyzed (audio required)',
                        'notes': ''
                    }
                    # For document-only input, use user-selected gender and default pitch analysis
                    results['pitch'] = {
                        'variation': 'Not analyzed (audio required)',
                        'consistency': 'Not analyzed (audio required)',
                        'average': 'Appropriate for the given gender (cannot be definitively determined without audio)',
                        'detected_gender': gender
                    }
                
                status_container.markdown("""
                    <div class="stStatusContainer stInfo">
                        <span class="step-counter">5</span>
                        Generating refined speech and detailed feedback...
                    </div>
                """, unsafe_allow_html=True)
                
                original, refined, _ = process_with_gpt(
                    openai_api_key, transcription, purpose, audience, duration, tone, additional_requirements, topic_of_speech, results
                )
                
                _, _, feedback = process_with_gemini(
                    transcription, purpose, audience, duration, tone, additional_requirements, topic_of_speech, results
                )

                status_container.markdown("""
                    <div class="stStatusContainer stInfo">
                        <span class="step-counter">6</span>
                        Generating and saving refined speech...
                    </div>
                """, unsafe_allow_html=True)
                
                text_filepath = save_processed_data(st.session_state.session_id, 'text', refined)
                st.session_state.text_filepath = text_filepath

                status_container.markdown("""
                    <div class="stStatusContainer stInfo">
                        <span class="step-counter">7</span>
                        Generating audio from refined speech...
                    </div>
                """, unsafe_allow_html=True)
                
                cleaned_speech = re.sub(r'\*\*|\*_pause_\*|_|#|<[^>]*>', '', refined)
                cleaned_speech = cleaned_speech.strip()
                
                if not cleaned_speech:
                    st.error("No text available for audio generation")
                    st.stop()

                tts = gTTS(text=cleaned_speech)
                audio_io = io.BytesIO()
                tts.write_to_fp(audio_io)
                audio_io.seek(0)
                audio_base64 = base64.b64encode(audio_io.read()).decode()

                if audio_base64:
                    audio_bytes = base64.b64decode(audio_base64)
                    audio_filepath = save_processed_data(st.session_state.session_id, 'audio', audio_bytes)
                    st.session_state.audio_filepath = audio_filepath

                status_container.markdown("""
                    <div class="stStatusContainer stSuccess">
                        <span class="step-counter">✓</span>
                        Processing complete! Showing results...
                    </div>
                """, unsafe_allow_html=True)

                # Store results in session state for persistence
                st.session_state.original = original
                st.session_state.refined = refined
                st.session_state.results = results
                st.session_state.pronunciation_data = pronunciation_data

                # Render the results after processing
                # content = f"""
                # <div class="content-section">
                #     <h3>Original Transcription</h3>
                #     <div class="transcription-container">
                #         <div class="transcription-legend">
                #             <strong>Legend:</strong><br>
                #             <span class="bold-word">Bold words</span> - Words to emphasize<br>
                #             <span class="pause-marker">|</span> - Pause in speech<br>
                #             <span class="mispronounced">Highlighted words</span> - Words that need pronunciation improvement
                #         </div>
                #         <div class="transcription-text">
                #             {format_transcription_text(original, [item['word'] for item in pronunciation_data['mispronounced_words']])}
                #         </div>
                #     </div>
                    
                #     <h3>Refined Speech</h3>
                #     {format_transcription_with_emphasis(refined, [item['word'] for item in pronunciation_data['mispronounced_words']])}
                # </div>
                # {format_detailed_feedback(results, pronunciation_data)}
                # """
                # st.html(content)

                # In the services function, before rendering the HTML content
                # st.write("Mispronounced Words Passed to Frontend:", [item['word'] for item in pronunciation_data['mispronounced_words']])

                content = f"""
                <div class="content-section">
                    <h3>Original Transcription</h3>
                    <div class="transcription-container">
                        <div class="transcription-legend">
                            <strong>Legend:</strong><br>
                            <span class="bold-word">Bold words</span> - Words to emphasize<br>
                            <span class="pause-marker">|</span> - Pause in speech<br>
                            <span class="mispronounced">Highlighted words</span> - Words that need pronunciation improvement
                        </div>
                        <div class="transcription-text">
                            {format_transcription_text(original, [item['word'] for item in pronunciation_data['mispronounced_words']])}
                        </div>
                    </div>
                    
                    <h3>Refined Speech</h3>
                    {format_transcription_with_emphasis(refined, [item['word'] for item in pronunciation_data['mispronounced_words']])}
                </div>
                {format_detailed_feedback(results, pronunciation_data)}
                """
                st.html(content)


                st.markdown("""
                <style>
                /* Existing CSS */
                .mispronounced {
                    background-color: rgba(255, 0, 0, 0.2);
                    color: #ff0000;
                    padding: 2px 4px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                .transcription-text {
                    white-space: pre-wrap;
                    word-wrap: break-word;
                    font-size: 16px;
                    line-height: 1.6;
                    padding: 15px;
                    background-color: #f8f9fa;
                    border-radius: 5px;
                }
                /* Rest of the CSS */
                </style>
            """, unsafe_allow_html=True)
                                
                # st.markdown("""
                #     <style>
                #         .mispronounced {
                #             background-color: rgba(255, 0, 0, 0.1);
                #             padding: 0 2px;
                #             border-radius: 2px;
                #             color: #ff0000;
                #             font-weight: bold;
                #         }
                #         .transcription-text {
                #             white-space: pre-wrap;
                #             word-wrap: break-word;
                #             font-size: 16px;
                #             line-height: 1.6;
                #             padding: 15px;
                #             background-color: #f8f9fa;
                #             border-radius: 5px;
                #         }
                #     </style>
                # """, unsafe_allow_html=True)
                
                if audio_base64:
                    content = """
                        <div class="generate-speech">
                            <h3>Generated Speech Audio</h3>
                        </div>
                        """
                    st.html(content)
                    st.subheader("AI Version of Refined Speech")
                    st.audio(io.BytesIO(audio_bytes), format="audio/wav")
                
                    if st.session_state.get('ai_audio_bytes'):
                        st.subheader("AI Version of Original Speech")
                        try:
                            ai_audio_io = io.BytesIO(st.session_state.ai_audio_bytes)
                            ai_audio_io.seek(0)
                            st.audio(ai_audio_io, format='audio/mp3')
                            
                            download_container = st.container()
                            
                            with download_container:
                                st.markdown(
                                    f'<a href="data:audio/mp3;base64,{base64.b64encode(st.session_state.ai_audio_bytes).decode()}" download="ai_version.mp3" style="text-decoration: none;">'
                                    '<button style="background-color: #4CAF50; color: white; padding: 10px 20px; border: none; border-radius: 4px; cursor: pointer; width: 100%;">'
                                    'Download AI Version'
                                    '</button>'
                                    '</a>',
                                    unsafe_allow_html=True
                                )
                            
                            with download_container:
                                st.markdown(
                                    f'<a href="data:audio/wav;base64,{audio_base64}" download="refined_speech.wav" style="text-decoration: none;">'
                                    '<button style="background-color: #4CAF50; color: white; padding: 10px 20px; border: none; border-radius: 4px; cursor: pointer; width: 100%;">'
                                    'Download Refined Speech Audio'
                                    '</button>'
                                    '</a>',
                                    unsafe_allow_html=True
                                )
                        except Exception as e:
                            st.error(f"Error displaying AI audio: {str(e)}")
                            st.write("Debug info - Audio data length:", len(st.session_state.ai_audio_bytes) if st.session_state.get('ai_audio_bytes') else "No audio data")
            
            except Exception as e:
                if not st.session_state.get('is_authenticated', False):
                    st.session_state.usage_count -= 1
                status_container.error(f"An error occurred: {str(e)}")
                st.stop()
            
    # if st.session_state.get('results'):
    #     st.subheader("Analysis Results")
        
    #     if uploaded_file:
    #         st.subheader("Original Speech")
    #         st.audio(uploaded_file, format='audio/wav')
            
    #         if st.session_state.get('ai_audio_bytes'):
    #             st.subheader("AI Version of Original Speech")
    #             try:
    #                 ai_audio_io = io.BytesIO(st.session_state.ai_audio_bytes)
    #                 ai_audio_io.seek(0)
    #                 st.audio(ai_audio_io, format='audio/mp3')
    #                 st.download_button(
    #                     label="Download AI Version",
    #                     data=st.session_state.ai_audio_bytes,
    #                     file_name="ai_version.mp3",
    #                     mime="audio/mp3"
    #                 )
    #             except Exception as e:
    #                 st.error(f"Error displaying AI audio: {str(e)}")

def main():
    st.set_page_config(
        page_title="SpeechSmith Services",
        page_icon="🎤",
        layout="wide"
    )
    services()

if __name__ == "__main__":
    main()


